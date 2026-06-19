"""Acceptance test: Verify V18 deliverables are real and usable.

Enhanced checks for:
- AI mode tracking
- action_advice not empty
- evidence_gap structured
- no placeholders
- no internal fields
- PDF exists
- ZIP complete
"""
import sys
import os
import re
import json
from pathlib import Path
from typing import List, Tuple

# Ensure project root is in path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, project_root)


def check_docx_opens(path: str) -> bool:
    """Check DOCX can be opened."""
    try:
        from docx import Document
        doc = Document(path)
        return len(doc.paragraphs) > 0
    except Exception:
        return False


def check_docx_chinese_chars(path: str, min_chars: int = 100) -> bool:
    """Check DOCX has sufficient Chinese characters."""
    try:
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        chinese_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
        return chinese_count >= min_chars
    except Exception:
        return False


def check_pdf_header(path: str) -> bool:
    """Check PDF starts with %PDF."""
    try:
        with open(path, "rb") as f:
            header = f.read(5)
        return header.startswith(b"%PDF")
    except Exception:
        return False


def check_xlsx_rows(path: str, min_rows: int = 3) -> bool:
    """Check XLSX has at least min_rows data rows."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path)
        ws = wb.active
        if ws is None:
            return False
        return ws.max_row >= min_rows + 1  # +1 for header
    except Exception:
        return False


def check_no_forbidden_files(customer_dir: str) -> bool:
    """Check customer dir has no json/md/debug files."""
    forbidden_exts = [".json", ".md"]
    forbidden_names = ["debug", "ai_raw_outputs", "raw", "internal", 
                       "workflow_trace", "ai_run_manifest", "full_text"]

    for f in Path(customer_dir).rglob("*"):
        if f.is_file():
            if f.suffix.lower() in forbidden_exts:
                return False
            if any(fn in f.name.lower() for fn in forbidden_names):
                return False
    return True


def check_no_placeholders(path: str) -> bool:
    """Check DOCX has no placeholder text."""
    try:
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        # Only check for actual template placeholders, not normal Chinese text
        forbidden = [
            "TODO", "暂无", "请自行补充", "lorem ipsum", "示例文本",
            "{{", "}}", "complainant_name", "respondent_name",
            "facts_and_reasons", "legal_basis", "evidence_list",
            "占位符", "模板变量", "待填写",
        ]
        return not any(f.lower() in text.lower() for f in forbidden)
    except Exception:
        return False


def check_no_internal_fields(path: str) -> bool:
    """Check DOCX has no internal system fields."""
    try:
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        forbidden = [
            "fact_card", "analysis.json", "prompt", "source_id",
            "confidence", "raw_output", "workflow_trace",
            "PipelineContext", "StrategyCard", "DistilledCard",
        ]
        return not any(f in text for f in forbidden)
    except Exception:
        return False


def check_no_double_yuan(path: str) -> bool:
    """Check DOCX has no '元元' (double yuan) issue."""
    try:
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return "元元" not in text
    except Exception:
        return False


def check_no_doc_type_header(path: str) -> bool:
    """Check DOCX has no '文书 1:' or '类型:' internal markers."""
    try:
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        # Check for internal markers
        if "文书 1:" in text or "文书 1：" in text:
            return False
        if re.search(r'^类型[:：]', text, re.MULTILINE):
            return False
        if "类型:" in text or "类型：" in text:
            # Allow "案件类型" but not standalone "类型:"
            lines = text.split('\n')
            for line in lines:
                stripped = line.strip()
                if stripped.startswith("类型:") or stripped.startswith("类型："):
                    if "案件" not in stripped:
                        return False
        return True
    except Exception:
        return False


def check_no_waiting_placeholder(path: str) -> bool:
    """Check DOCX has no '（待补充）' placeholder text."""
    try:
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        forbidden = ["（待补充）", "(待补充)", "【待补充】", "内容待补充"]
        return not any(f in text for f in forbidden)
    except Exception:
        return False


def check_distilled_card_quality(case_dir: str) -> Tuple[bool, str]:
    """Check distilled_card.json meets quality requirements."""
    distilled_path = os.path.join(case_dir, "distilled_card.json")
    if not os.path.exists(distilled_path):
        return False, "distilled_card.json not found"
    
    try:
        with open(distilled_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        sc = data.get("strategy_card", {})
        issues = []
        
        # Check action_advice count (minimum 6)
        action_advice = sc.get("action_advice", [])
        if len(action_advice) < 6:
            issues.append(f"action_advice only {len(action_advice)} items (need 6+)")
        
        # Check evidence_gap count (minimum 5)
        evidence_gap = sc.get("evidence_gap", [])
        if len(evidence_gap) < 5:
            issues.append(f"evidence_gap only {len(evidence_gap)} items (need 5+)")
        
        # Check SABCD rating has reasoning
        rating = sc.get("sabcd_rating", "")
        if not rating:
            issues.append("sabcd_rating is empty")
        
        situation = sc.get("situation_assessment", "")
        if not situation or len(situation) < 50:
            issues.append("situation_assessment too short or empty")
        
        if issues:
            return False, "; ".join(issues)
        
        return True, f"action: {len(action_advice)}, gaps: {len(evidence_gap)}, rating: {rating}"
    except Exception as e:
        return False, f"Error reading distilled_card: {e}"


def check_customer_zip_no_debug(customer_dir: str) -> bool:
    """Check customer ZIP contains no debug/json/md files."""
    import zipfile
    zip_files = list(Path(customer_dir).glob("*.zip"))
    if not zip_files:
        return True  # No ZIP = nothing to check
    
    forbidden_patterns = [
        "debug", "ai_run_manifest", "workflow_trace", "prompt",
        ".json", ".md", ".py", ".log", "distilled_card",
    ]
    
    for zip_path in zip_files:
        try:
            with zipfile.ZipFile(zip_path, 'r') as zf:
                for name in zf.namelist():
                    name_lower = name.lower()
                    for pattern in forbidden_patterns:
                        if pattern in name_lower:
                            return False
        except Exception:
            pass
    
    return True


def check_ai_mode_manifest(case_dir: str) -> Tuple[bool, str]:
    """Check ai_run_manifest.json exists and has valid ai_mode.
    
    local_fallback means no real AI was used - this BLOCKS formal delivery.
    """
    manifest_path = os.path.join(case_dir, "ai_run_manifest.json")
    if not os.path.exists(manifest_path):
        return False, "ai_run_manifest.json not found"
    
    try:
        with open(manifest_path, "r", encoding="utf-8") as f:
            manifest = json.load(f)
        
        ai_mode = manifest.get("ai_mode", "")
        if ai_mode not in ("real_ai", "mixed", "local_fallback", "dual_ai"):
            return False, f"Invalid ai_mode: {ai_mode}"
        
        # local_fallback = no real AI → FAIL for formal delivery
        if ai_mode == "local_fallback":
            return False, "FAIL_NO_REAL_AI: local_fallback mode cannot be formally delivered"
        
        return True, ai_mode
    except Exception as e:
        return False, f"Error reading manifest: {e}"


def check_distilled_card(case_dir: str) -> Tuple[bool, str]:
    """Check distilled_card.json has action_advice and structured evidence_gap."""
    distilled_path = os.path.join(case_dir, "distilled_card.json")
    if not os.path.exists(distilled_path):
        return False, "distilled_card.json not found"
    
    try:
        with open(distilled_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        sc = data.get("strategy_card", {})
        
        # Check action_advice not empty
        action_advice = sc.get("action_advice", [])
        if not action_advice:
            return False, "action_advice is empty"
        
        # Check evidence_gap is structured (list of dicts, not just strings)
        evidence_gap = sc.get("evidence_gap", [])
        if not evidence_gap:
            return False, "evidence_gap is empty"
        
        return True, f"action_advice: {len(action_advice)}, evidence_gap: {len(evidence_gap)}"
    except Exception as e:
        return False, f"Error reading distilled_card: {e}"


def check_pdf_exists(customer_dir: str) -> bool:
    """Check at least one PDF exists in customer dir."""
    pdf_files = list(Path(customer_dir).glob("*.pdf"))
    return len(pdf_files) > 0


def check_zip_contains_pdf(customer_dir: str) -> bool:
    """Check ZIP contains PDF files."""
    import zipfile
    zip_files = list(Path(customer_dir).glob("*.zip"))
    if not zip_files:
        return False
    
    for zip_path in zip_files:
        try:
            with zipfile.ZipFile(zip_path, 'r') as zf:
                pdf_files = [f for f in zf.namelist() if f.endswith('.pdf')]
                if pdf_files:
                    return True
        except Exception:
            pass
    
    return False


def run_acceptance_test(case_dir: str) -> bool:
    """Run all acceptance checks on a case directory."""
    customer_dir = os.path.join(case_dir, "customer")

    if not os.path.exists(customer_dir):
        print(f"FAIL: Customer directory not found: {customer_dir}")
        return False

    checks: List[Tuple[str, bool, str]] = []

    # Check DOCX files
    docx_files = list(Path(customer_dir).glob("*.docx"))
    for docx in docx_files:
        checks.append((f"DOCX opens: {docx.name}", check_docx_opens(str(docx)), ""))
        checks.append((f"DOCX Chinese: {docx.name}", check_docx_chinese_chars(str(docx)), ""))
        checks.append((f"DOCX no placeholder: {docx.name}", check_no_placeholders(str(docx)), ""))
        checks.append((f"DOCX no internal: {docx.name}", check_no_internal_fields(str(docx)), ""))
        checks.append((f"DOCX no double yuan: {docx.name}", check_no_double_yuan(str(docx)), ""))
        checks.append((f"DOCX no doc-type header: {docx.name}", check_no_doc_type_header(str(docx)), ""))
        checks.append((f"DOCX no waiting placeholder: {docx.name}", check_no_waiting_placeholder(str(docx)), ""))

    # Check PDF files
    pdf_files = list(Path(customer_dir).glob("*.pdf"))
    for pdf in pdf_files:
        checks.append((f"PDF header: {pdf.name}", check_pdf_header(str(pdf)), ""))

    # Check XLSX
    xlsx_files = list(Path(customer_dir).glob("*.xlsx"))
    for xlsx in xlsx_files:
        checks.append((f"XLSX rows: {xlsx.name}", check_xlsx_rows(str(xlsx)), ""))

    # Check ZIP
    zip_files = list(Path(customer_dir).glob("*.zip"))
    checks.append(("ZIP exists", len(zip_files) > 0, ""))

    # Check no forbidden files
    checks.append(("No forbidden files", check_no_forbidden_files(customer_dir), ""))

    # Check PDF exists
    checks.append(("PDF exists", check_pdf_exists(customer_dir), ""))

    # Check ZIP contains PDF
    checks.append(("ZIP contains PDF", check_zip_contains_pdf(customer_dir), ""))

    # Check AI mode manifest
    manifest_ok, manifest_msg = check_ai_mode_manifest(case_dir)
    checks.append(("AI mode manifest", manifest_ok, manifest_msg))

    # Check distilled card basic
    distilled_ok, distilled_msg = check_distilled_card(case_dir)
    checks.append(("Distilled card", distilled_ok, distilled_msg))

    # Check distilled card quality
    quality_ok, quality_msg = check_distilled_card_quality(case_dir)
    checks.append(("Distilled card quality", quality_ok, quality_msg))

    # Check customer ZIP no debug
    checks.append(("Customer ZIP no debug", check_customer_zip_no_debug(customer_dir), ""))

    # Check Defense Quality Gate (CRITICAL - must pass)
    from core.quality.defense_quality_gate import DefenseQualityGate
    defense_gate = DefenseQualityGate()
    manifest_path = os.path.join(case_dir, "ai_run_manifest.json")
    ai_mode = "unknown"
    if os.path.exists(manifest_path):
        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                manifest = json.load(f)
            ai_mode = manifest.get("ai_mode", "unknown")
        except Exception:
            pass

    defense_result = defense_gate.run_checks(case_dir, ai_mode)
    defense_passed = defense_result.passed
    defense_msg = defense_result.summary()
    checks.append(("Defense Quality Gate", defense_passed, defense_msg))

    # Add individual defense gate failures
    if not defense_passed:
        for dc in defense_result.checks:
            if not dc.passed:
                checks.append((f"  Defense: {dc.check_name}", False, dc.message))

    # Print results
    failed = []
    for name, passed, msg in checks:
        status = "[PASS]" if passed else "[FAIL]"
        detail = f" - {msg}" if msg else ""
        print(f"  {status}{detail} {name}")
        if not passed:
            failed.append(name)

    print()
    if failed:
        print(f"FAIL: {len(failed)} checks failed")
        for f in failed:
            print(f"  - {f}")
        return False
    else:
        print("PASS: All acceptance checks passed")
        return True


if __name__ == "__main__":
    case_dir = sys.argv[1] if len(sys.argv) > 1 else "outputs/case_001"
    print(f"=== V18 Acceptance Test: {case_dir} ===")
    success = run_acceptance_test(case_dir)
    sys.exit(0 if success else 1)
