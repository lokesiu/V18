"""verify_v3_output.py — 验证 v3 输出文档内容 + DB 状态"""
import sys
sys.path.insert(0, '.')
sys.dont_write_bytecode = True

import os, re, json
from pathlib import Path

OUT_DIR = r"D:\codex\V18\test_run_judgment\output_v3"

# 1. 验证再审申请书 DOCX 内容
docx_path = os.path.join(OUT_DIR, "customer", "06_再审申请书_江宁区全至惠百货超市店案.docx")
pdf_path = os.path.join(OUT_DIR, "customer", "06_再审申请书_江宁区全至惠百货超市店案.pdf")

print("=" * 70)
print(f"DOCX 存在: {os.path.exists(docx_path)} ({os.path.getsize(docx_path) if os.path.exists(docx_path) else 0} bytes)")
print(f"PDF 存在: {os.path.exists(pdf_path)} ({os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 0} bytes)")
print()

# 用 python-docx 读 DOCX 文本
try:
    from docx import Document
    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    print(f"DOCX 段落数: {len(doc.paragraphs)}")
    print(f"DOCX 总字符数: {len(full_text)}")
    print()
    print("=" * 70)
    print("【禁用词扫描】")
    print("=" * 70)
    from core.pipeline.external_doc_filter import FORBIDDEN_PHRASES
    banned = set(p[0] for p in FORBIDDEN_PHRASES)
    found_violations = []
    for phrase in banned:
        if phrase in full_text:
            found_violations.append(phrase)
    if found_violations:
        print(f"❌ 发现 {len(found_violations)} 处禁用词:")
        for v in found_violations:
            line_no = None
            for i, line in enumerate(full_text.split('\n'), 1):
                if v in line:
                    line_no = i
                    break
            print(f"  - 第 {line_no} 行「{v}」")
    else:
        print(f"✅ 禁用词扫描通过(扫描 {len(banned)} 个词,0 处违规)")

    print()
    print("=" * 70)
    print("【关键章节检查】")
    print("=" * 70)
    bad_sections = ["风险与再审必要性", "风险评估", "风险提示", "败诉风险", "成本分析"]
    bad_section_found = False
    for sec in bad_sections:
        if sec in full_text:
            bad_section_found = True
            print(f"  ❌ 含非法章节「{sec}」")
    if not bad_section_found:
        print(f"✅ 无非法章节(扫描 {len(bad_sections)} 个常见内部章节名)")

    print()
    print("=" * 70)
    print("【当事人信息检查】")
    print("=" * 70)
    checks = [
        ("郑恢快", "原告姓名"),
        ("35042519920627071X", "原告身份证号"),
        ("江宁区全至惠百货超市店", "店名"),
        ("92320115MACGE2GL9N", "统一社会信用代码"),
        ("刘宝旺", "经营者"),
        ("2026年", "年份(2026 应一致)"),
        ("江苏省南京市中级人民法院", "再审法院"),
    ]
    for kw, label in checks:
        present = "✅" if kw in full_text else "❌"
        print(f"  {present} {label}: '{kw}' {'在' if kw in full_text else '不在'}文中")

    print()
    print("=" * 70)
    print("【全文预览(前 1500 字)】")
    print("=" * 70)
    print(full_text[:1500])
    print("...")
    print(full_text[-1000:] if len(full_text) > 2500 else "")
except Exception as exc:
    print(f"DOCX 读取失败: {exc}")
    import traceback
    traceback.print_exc()

# 2. 验证 DB 状态
print()
print("=" * 70)
print("【DB 人审状态】")
print("=" * 70)
try:
    from core.task_store import get_task_store
    store = get_task_store()
    # 列出最近的任务
    tasks = store.list_recent(limit=5)
    if tasks:
        for t in tasks:
            print(f"任务: {t.task_id[:25]}... | identity={t.identity} | goal={t.goal}")
            print(f"  status={t.status}, human_review_status={t.human_review_status}")
            print(f"  human_review_pending_count={t.human_review_pending_count}")
            docs = store.get_human_review_docs(t.task_id)
            for d in docs:
                print(f"    - {d.get('doc_type')}: {d.get('status')}")
                print(f"      source={d.get('source')}, warnings={len(d.get('warnings', []))}")
                print(f"      sections_removed={len(d.get('sections_removed', []))}")
            print()
except Exception as exc:
    print(f"DB 查询失败: {exc}")
    import traceback
    traceback.print_exc()
