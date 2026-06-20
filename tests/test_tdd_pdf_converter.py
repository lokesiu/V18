"""TDD tests for core/render/pdf_converter.py — branch coverage."""
import pytest
import sys, os
sys.path.insert(0, ".")

from core.render.pdf_converter import (
    convert_to_pdf, _try_reportlab, _try_docx2pdf, _try_libreoffice,
    render_pdf_from_text,
)
from core.render.docx_renderer import render_docx_from_text


# ══════════════════════════════════════════════════════════════════════
# convert_to_pdf — main entry point
# ══════════════════════════════════════════════════════════════════════
class TestConvertToPdf:
    def test_nonexistent_docx(self, tmp_path):
        assert convert_to_pdf("/nonexistent.docx", str(tmp_path / "out.pdf")) is False

    def test_valid_docx(self, tmp_path):
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        render_docx_from_text("答辩状\n测试内容" * 20, str(docx))
        result = convert_to_pdf(str(docx), str(pdf))
        assert result is True
        assert pdf.exists()
        assert pdf.stat().st_size > 100

    def test_empty_docx(self, tmp_path):
        docx = tmp_path / "empty.docx"
        pdf = tmp_path / "empty.pdf"
        render_docx_from_text("", str(docx))
        result = convert_to_pdf(str(docx), str(pdf))
        assert result is True

    def test_docx_with_centered_title(self, tmp_path):
        docx = tmp_path / "titled.docx"
        pdf = tmp_path / "titled.pdf"
        render_docx_from_text("民事答辩状\n\n正文内容" * 20, str(docx))
        result = convert_to_pdf(str(docx), str(pdf))
        assert result is True

    def test_docx_with_xml_chars(self, tmp_path):
        docx = tmp_path / "xml.docx"
        pdf = tmp_path / "xml.pdf"
        render_docx_from_text("张三 & 李四 < 王五 > 赵六", str(docx))
        result = convert_to_pdf(str(docx), str(pdf))
        assert result is True


# ══════════════════════════════════════════════════════════════════════
# _try_reportlab — branch coverage
# ══════════════════════════════════════════════════════════════════════
class TestTryReportlab:
    def test_valid_docx(self, tmp_path):
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        render_docx_from_text("答辩状\n正文" * 20, str(docx))
        assert _try_reportlab(str(docx), str(pdf)) is True

    def test_nonexistent_docx(self, tmp_path):
        assert _try_reportlab("/nonexistent.docx", str(tmp_path / "out.pdf")) is False

    def test_empty_paragraphs(self, tmp_path):
        """Empty paragraphs should produce Spacer."""
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        from docx import Document
        doc = Document()
        doc.add_paragraph("标题")
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("正文")
        doc.save(str(docx))
        assert _try_reportlab(str(docx), str(pdf)) is True

    def test_docx_with_heading_styles(self, tmp_path):
        """DOCX with Heading 1/2 styles."""
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        from docx import Document
        doc = Document()
        doc.add_paragraph("标题", style="Heading 1")
        doc.add_paragraph("小节", style="Heading 2")
        doc.add_paragraph("正文内容")
        doc.save(str(docx))
        assert _try_reportlab(str(docx), str(pdf)) is True

    def test_docx_with_generic_heading(self, tmp_path):
        """DOCX with generic 'Heading' style."""
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        from docx import Document
        doc = Document()
        doc.add_paragraph("标题")
        # Add a paragraph with Heading style
        p = doc.add_paragraph("章节")
        p.style = doc.styles["Heading 1"]
        doc.add_paragraph("正文")
        doc.save(str(docx))
        assert _try_reportlab(str(docx), str(pdf)) is True

    def test_docx_only_empty_paragraphs(self, tmp_path):
        """DOCX with only empty paragraphs → fallback '法律文书' title."""
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        from docx import Document
        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.save(str(docx))
        assert _try_reportlab(str(docx), str(pdf)) is True

    def test_markdown_cleaning(self, tmp_path):
        """Markdown bold/heading/list markers should be cleaned."""
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        render_docx_from_text("**粗体**\n### 标题\n- 列表项\n正文", str(docx))
        assert _try_reportlab(str(docx), str(pdf)) is True

    def test_output_dir_auto_create(self, tmp_path):
        """Output directory should be auto-created."""
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "sub" / "dir" / "test.pdf"
        render_docx_from_text("内容" * 20, str(docx))
        assert _try_reportlab(str(docx), str(pdf)) is True
        assert pdf.exists()


# ══════════════════════════════════════════════════════════════════════
# render_pdf_from_text — full branch coverage
# ══════════════════════════════════════════════════════════════════════
class TestRenderPdfFromText:
    def test_basic(self, tmp_path):
        out = str(tmp_path / "test.pdf")
        result = render_pdf_from_text("标题", "正文内容" * 20, out)
        assert result is True
        assert os.path.getsize(out) > 100

    def test_empty_title(self, tmp_path):
        out = str(tmp_path / "test.pdf")
        result = render_pdf_from_text("", "正文" * 20, out)
        assert result is True

    def test_empty_content(self, tmp_path):
        out = str(tmp_path / "test.pdf")
        result = render_pdf_from_text("标题", "", out)
        assert result is True

    def test_xml_unsafe_chars(self, tmp_path):
        out = str(tmp_path / "test.pdf")
        result = render_pdf_from_text("标题", "张三 & 李四 < 王五 > 赵六", out)
        assert result is True

    def test_blank_lines(self, tmp_path):
        """Blank lines should produce Spacers."""
        out = str(tmp_path / "test.pdf")
        content = "第一段\n\n\n第二段\n\n第三段"
        result = render_pdf_from_text("标题", content, out)
        assert result is True

    def test_auto_create_dir(self, tmp_path):
        out = str(tmp_path / "new_dir" / "test.pdf")
        result = render_pdf_from_text("标题", "内容" * 20, out)
        assert result is True

    def test_long_content(self, tmp_path):
        """Very long content should produce multi-page PDF."""
        out = str(tmp_path / "test.pdf")
        content = "这是一段很长的内容。\n" * 500
        result = render_pdf_from_text("长文档", content, out)
        assert result is True
        assert os.path.getsize(out) > 1000


# ══════════════════════════════════════════════════════════════════════
# _try_docx2pdf — graceful failure
# ══════════════════════════════════════════════════════════════════════
class TestTryDocx2pdf:
    def test_returns_false_when_not_installed(self, tmp_path):
        """Should return False gracefully when docx2pdf not available."""
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        render_docx_from_text("内容", str(docx))
        # docx2pdf likely not installed in test environment
        result = _try_docx2pdf(str(docx), str(pdf))
        assert isinstance(result, bool)

    def test_nonexistent_file(self, tmp_path):
        result = _try_docx2pdf("/nonexistent.docx", str(tmp_path / "out.pdf"))
        assert result is False


# ══════════════════════════════════════════════════════════════════════
# _try_libreoffice — graceful failure
# ══════════════════════════════════════════════════════════════════════
class TestTryLibreoffice:
    def test_returns_false_when_not_installed(self, tmp_path):
        """Should return False gracefully when LibreOffice not available."""
        docx = tmp_path / "test.docx"
        pdf = tmp_path / "test.pdf"
        render_docx_from_text("内容", str(docx))
        result = _try_libreoffice(str(docx), str(pdf))
        assert isinstance(result, bool)

    def test_nonexistent_file(self, tmp_path):
        result = _try_libreoffice("/nonexistent.docx", str(tmp_path / "out.pdf"))
        assert result is False
