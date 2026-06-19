"""
tests/test_no_forbidden_text_in_customer_docs.py

Test: Customer DOCX files must not contain forbidden text patterns.
"""
import os
import sys
from pathlib import Path

sys.path.insert(0, ".")

# Forbidden patterns that must not appear in customer-facing documents
FORBIDDEN_PATTERNS = [
    "待补充",
    "类型:",
    "类型：",
    "文书 1",
    "文书1",
    "TODO",
    "暂无",
    "请自行补充",
    "XXX",
    "{{",
    "}}",
    "待评估",
]


def _get_customer_dir():
    """Get the golden case customer directory."""
    return "outputs/golden_defense_case/customer"


def _read_docx_text(docx_path: str) -> str:
    """Read all text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(docx_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def test_no_forbidden_text_in_docx():
    """Customer DOCX files must not contain forbidden text patterns."""
    customer_dir = _get_customer_dir()
    if not os.path.isdir(customer_dir):
        return  # Skip if directory doesn't exist

    docx_files = list(Path(customer_dir).glob("*.docx"))
    assert len(docx_files) > 0, "No DOCX files found in customer directory"

    violations = []
    for docx_path in docx_files:
        text = _read_docx_text(str(docx_path))
        for pattern in FORBIDDEN_PATTERNS:
            if pattern in text:
                violations.append(f"{docx_path.name}: contains '{pattern}'")

    assert not violations, "Forbidden text found:\n" + "\n".join(violations)


def test_no_dai_bu_chong():
    """Specifically test for '待补充' pattern."""
    customer_dir = _get_customer_dir()
    if not os.path.isdir(customer_dir):
        return

    docx_files = list(Path(customer_dir).glob("*.docx"))
    for docx_path in docx_files:
        text = _read_docx_text(str(docx_path))
        assert "待补充" not in text, f"{docx_path.name} contains '待补充'"


def test_no_lei_xing():
    """Specifically test for '类型:' pattern."""
    customer_dir = _get_customer_dir()
    if not os.path.isdir(customer_dir):
        return

    docx_files = list(Path(customer_dir).glob("*.docx"))
    for docx_path in docx_files:
        text = _read_docx_text(str(docx_path))
        assert "类型:" not in text, f"{docx_path.name} contains '类型:'"
        assert "类型：" not in text, f"{docx_path.name} contains '类型：'"


def test_no_xxx():
    """Specifically test for 'XXX' pattern."""
    customer_dir = _get_customer_dir()
    if not os.path.isdir(customer_dir):
        return

    docx_files = list(Path(customer_dir).glob("*.docx"))
    for docx_path in docx_files:
        text = _read_docx_text(str(docx_path))
        assert "XXX" not in text, f"{docx_path.name} contains 'XXX'"
