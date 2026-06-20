"""TDD tests for step7_render.py and step7_render_manifest.py pure functions."""
import pytest
import sys, os, json
sys.path.insert(0, ".")

from core.fact_card import (
    PipelineContext, FactCard, Party, SourceRef,
    StrategyCard, ActionAdvice, DraftDocument, DistilledCard,
)
from core.pipeline.step7_render import (
    _get_identity_extra_doc, _derive_case_type,
    _extract_user_info_from_refs, _get_party_name,
    _make_filename, _render_docx_from_content,
    _render_xlsx_from_fact_card, _render_strategy_docx,
    _render_draft_documents, _manual_build_zip,
    _save_distilled_card_as_text,
)
from core.pipeline.step7_render_manifest import (
    is_retryable, _file_size, _check_source_dependency,
    render_with_manifest,
)


# ══════════════════════════════════════════════════════════════════════
# _get_identity_extra_doc
# ══════════════════════════════════════════════════════════════════════
class TestGetIdentityExtraDoc:
    def test_defendant_full(self):
        r = _get_identity_extra_doc("被诉方（被告）")
        assert r is not None
        assert r[1] == "答辩状"
        assert r[2] == "docx"

    def test_defendant_brief(self):
        r = _get_identity_extra_doc("被诉方")
        assert r is not None
        assert r[1] == "答辩状"

    def test_plaintiff(self):
        r = _get_identity_extra_doc("起诉方")
        assert r[1] == "起诉状"

    def test_complainant(self):
        r = _get_identity_extra_doc("投诉方")
        assert r[1] == "投诉状"

    def test_review_applicant(self):
        r = _get_identity_extra_doc("行政复议申请人")
        assert r[1] == "行政复议申请书"

    def test_unknown_returns_none(self):
        assert _get_identity_extra_doc("未知身份") is None
        assert _get_identity_extra_doc("") is None
        assert _get_identity_extra_doc("整理证据") is None


# ══════════════════════════════════════════════════════════════════════
# _derive_case_type
# ══════════════════════════════════════════════════════════════════════
class TestDeriveCaseType:
    def test_none_fc(self):
        assert _derive_case_type(None) == "待补充"

    def test_empty_fc(self):
        assert _derive_case_type(FactCard()) == "民事纠纷"

    def test_loan_keyword(self):
        fc = FactCard(key_facts=["被告借款10万元"])
        assert _derive_case_type(fc) == "民间借贷纠纷"

    def test_loan_lending(self):
        fc = FactCard(key_facts=["借贷关系"])
        assert _derive_case_type(fc) == "民间借贷纠纷"

    def test_loan_advance(self):
        fc = FactCard(key_facts=["垫资协议"])
        assert _derive_case_type(fc) == "民间借贷纠纷"

    def test_sale(self):
        fc = FactCard(key_facts=["买卖合同"])
        assert _derive_case_type(fc) == "买卖合同纠纷"

    def test_rental(self):
        fc = FactCard(key_facts=["租赁房屋"])
        assert _derive_case_type(fc) == "租赁合同纠纷"

    def test_labor(self):
        fc = FactCard(key_facts=["劳动争议"])
        assert _derive_case_type(fc) == "劳动争议"

    def test_injury(self):
        fc = FactCard(key_facts=["工伤认定"])
        assert _derive_case_type(fc) == "劳动争议"

    def test_contract(self):
        fc = FactCard(key_facts=["合同违约"])
        assert _derive_case_type(fc) == "合同纠纷"

    def test_agreement(self):
        fc = FactCard(key_facts=["协议纠纷"])
        assert _derive_case_type(fc) == "合同纠纷"

    def test_tort(self):
        fc = FactCard(key_facts=["侵权行为"])
        assert _derive_case_type(fc) == "侵权责任纠纷"

    def test_agency(self):
        fc = FactCard(key_facts=["中介服务"])
        assert _derive_case_type(fc) == "中介合同纠纷"

    def test_priority_order(self):
        """借款 takes priority over 合同."""
        fc = FactCard(key_facts=["借款合同"])
        assert _derive_case_type(fc) == "民间借贷纠纷"

    def test_no_facts(self):
        fc = FactCard(key_facts=None)
        assert _derive_case_type(fc) == "民事纠纷"


# ══════════════════════════════════════════════════════════════════════
# _extract_user_info_from_refs
# ══════════════════════════════════════════════════════════════════════
class TestExtractUserInfoFromRefs:
    def test_none_fc(self):
        assert _extract_user_info_from_refs(None) == {}

    def test_no_source_refs(self):
        fc = FactCard(parties=[Party(name="张三", role="被告")])
        assert _extract_user_info_from_refs(fc) == {}

    def test_no_parties(self):
        fc = FactCard(source_refs=[SourceRef(excerpt="test")])
        assert _extract_user_info_from_refs(fc) == {}

    def test_no_defendants(self):
        fc = FactCard(
            parties=[Party(name="张三", role="原告")],
            source_refs=[SourceRef(excerpt="申请人：张三")],
        )
        assert _extract_user_info_from_refs(fc) == {}

    def test_defendant_not_filed_by(self):
        """Source merely mentions defendant, not filed BY them."""
        fc = FactCard(
            parties=[Party(name="张三", role="被告")],
            source_refs=[SourceRef(excerpt="被告张三被起诉")],
        )
        assert _extract_user_info_from_refs(fc) == {}

    def test_defendant_applicant(self):
        fc = FactCard(
            parties=[Party(name="张三", role="被告")],
            source_refs=[SourceRef(excerpt="申请人：张三，男，1990年1月15日出生，住北京市朝阳区，身份证号：11010119900115001X，电话：13800138000")],
        )
        info = _extract_user_info_from_refs(fc)
        assert "张三" in info
        assert info["张三"]["gender"] == "男"
        assert info["张三"]["birth"] == "1990年1月15日"
        assert info["张三"]["id_number"] == "11010119900115001X"
        assert info["张三"]["phone"] == "13800138000"

    def test_defendant_respondent(self):
        fc = FactCard(
            parties=[Party(name="李四", role="被告")],
            source_refs=[SourceRef(excerpt="答辩人：李四，女，住上海市浦东新区")],
        )
        info = _extract_user_info_from_refs(fc)
        assert "李四" in info
        assert info["李四"]["gender"] == "女"

    def test_defendant_submitter(self):
        fc = FactCard(
            parties=[Party(name="王五", role="被告")],
            source_refs=[SourceRef(excerpt="提交人：王五，男")],
        )
        info = _extract_user_info_from_refs(fc)
        assert "王五" in info

    def test_multiple_defendants(self):
        fc = FactCard(
            parties=[
                Party(name="张三", role="被告"),
                Party(name="李四", role="被告"),
            ],
            source_refs=[
                SourceRef(excerpt="申请人：张三，男"),
                SourceRef(excerpt="答辩人：李四，女"),
            ],
        )
        info = _extract_user_info_from_refs(fc)
        assert "张三" in info
        assert "李四" in info

    def test_no_matching_defendant_in_excerpt(self):
        fc = FactCard(
            parties=[Party(name="张三", role="被告")],
            source_refs=[SourceRef(excerpt="申请人：张三，男")],
        )
        info = _extract_user_info_from_refs(fc)
        assert "张三" in info


# ══════════════════════════════════════════════════════════════════════
# _get_party_name
# ══════════════════════════════════════════════════════════════════════
class TestGetPartyName:
    def test_none_fact_card(self):
        ctx = PipelineContext()
        assert _get_party_name(ctx) == ""

    def test_no_parties(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard()
        assert _get_party_name(ctx) == ""

    def test_defendant_found(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard(parties=[
            Party(name="原告名", role="原告"),
            Party(name="被告名", role="被告"),
        ])
        assert _get_party_name(ctx) == "被告名"

    def test_no_defendant_use_plaintiff(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard(parties=[
            Party(name="原告名", role="原告"),
        ])
        assert _get_party_name(ctx) == "原告名"

    def test_no_plaintiff_use_first(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard(parties=[
            Party(name="第三人", role="第三人"),
        ])
        assert _get_party_name(ctx) == "第三人"

    def test_first_party_empty_name(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard(parties=[
            Party(name="", role="被告"),
        ])
        assert _get_party_name(ctx) == ""

    def test_multiple_defendants_first(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard(parties=[
            Party(name="被告A", role="被告"),
            Party(name="被告B", role="被告"),
        ])
        assert _get_party_name(ctx) == "被告A"


# ══════════════════════════════════════════════════════════════════════
# _make_filename
# ══════════════════════════════════════════════════════════════════════
class TestMakeFilename:
    def test_with_party(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard(parties=[Party(name="张三", role="被告")])
        fn = _make_filename("01_报告", "docx", ctx)
        assert fn == "01_报告_张三案.docx"

    def test_no_party(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard()
        fn = _make_filename("01_报告", "docx", ctx)
        assert fn == "01_报告.docx"

    def test_with_doc_type(self):
        ctx = PipelineContext()
        ctx.fact_card = FactCard(parties=[Party(name="李四", role="被告")])
        fn = _make_filename("06", "docx", ctx, doc_type="答辩状")
        assert fn == "06_答辩状_李四案.docx"

    def test_xlsx_format(self):
        ctx = PipelineContext()
        fn = _make_filename("04_证据目录", "xlsx", ctx)
        assert fn.endswith(".xlsx")

    def test_pdf_format(self):
        ctx = PipelineContext()
        fn = _make_filename("01_报告", "pdf", ctx)
        assert fn.endswith(".pdf")


# ══════════════════════════════════════════════════════════════════════
# _render_docx_from_content
# ══════════════════════════════════════════════════════════════════════
class TestRenderDocxFromContent:
    def test_empty_content(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        assert _render_docx_from_content("", out, ctx) is True
        assert os.path.getsize(out) > 0

    def test_markdown_bold_removed(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("**粗体**文本", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "**" not in text
        assert "粗体" in text

    def test_markdown_heading_removed(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("### 标题", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "###" not in text
        assert "标题" in text

    def test_xxx_replaced(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("XXX公司", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "XXX" not in text
        assert "已脱敏" in text

    def test_bracket_placeholder_removed(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("内容[请填写金额]继续", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "请填写" not in text

    def test_type_replaced(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("类型：证据", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "分类：" in text

    def test_id_preserved_in_info_line(self, tmp_path):
        """身份证号 in info lines should NOT be masked."""
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("身份证号：11010119900101001X", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "11010119900101001X" in text

    def test_id_masked_in_other_lines(self, tmp_path):
        """ID numbers in non-info lines should be masked."""
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("他的号码是11010119900101001X没错", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "11010119900101001X" not in text

    def test_phone_preserved_in_info_line(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("电话：13800138000", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "13800138000" in text

    def test_llm_opening_removed(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        _render_docx_from_content("好的，以下是答辩状。\n\n正文内容", out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "好的" not in text


# ══════════════════════════════════════════════════════════════════════
# _render_xlsx_from_fact_card
# ══════════════════════════════════════════════════════════════════════
class TestRenderXlsxFromFactCard:
    def test_basic(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.xlsx")
        fc = FactCard(case_id="C1")
        assert _render_xlsx_from_fact_card(fc, out, ctx) is True
        assert os.path.getsize(out) > 0

    def test_with_refs(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.xlsx")
        fc = FactCard(
            case_id="C1",
            source_refs=[SourceRef(file_name="a.pdf", page=1, excerpt="证据")],
        )
        assert _render_xlsx_from_fact_card(fc, out, ctx) is True


# ══════════════════════════════════════════════════════════════════════
# _render_strategy_docx
# ══════════════════════════════════════════════════════════════════════
class TestRenderStrategyDocx:
    def test_assessment(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        sc = StrategyCard(sabcd_rating="B", situation_assessment="评估内容")
        assert _render_strategy_docx(sc, "案件处境评估报告", out, ctx) is True

    def test_action_advice(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        sc = StrategyCard(action_advice=[ActionAdvice(action="做X", priority="S", reasoning="因为Y")])
        assert _render_strategy_docx(sc, "行动建议书", out, ctx) is True

    def test_evidence_gap_with_gaps(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        sc = StrategyCard(evidence_gap=["缺口A", "缺口B"])
        assert _render_strategy_docx(sc, "证据闭环补强清单", out, ctx) is True

    def test_evidence_gap_no_gaps(self, tmp_path):
        ctx = PipelineContext()
        out = str(tmp_path / "test.docx")
        sc = StrategyCard(evidence_gap=[])
        assert _render_strategy_docx(sc, "证据闭环补强清单", out, ctx) is True

    def test_unknown_doc_type(self, tmp_path):
        ctx = PipelineContext()
        ctx.identity = "被诉方"
        ctx.goal = "应诉答辩"
        out = str(tmp_path / "test.docx")
        sc = StrategyCard()
        assert _render_strategy_docx(sc, "未知类型", out, ctx) is True


# ══════════════════════════════════════════════════════════════════════
# _render_draft_documents
# ══════════════════════════════════════════════════════════════════════
class TestRenderDraftDocuments:
    def test_defendant_generates_defense(self, tmp_path):
        ctx = PipelineContext()
        ctx.identity = "被诉方（被告）"
        ctx.goal = "应诉答辩"
        ctx.fact_card = FactCard(
            case_id="C1",
            parties=[Party(name="张三", role="被告"), Party(name="李四", role="原告")],
            key_facts=["事实1"],
        )
        out = str(tmp_path / "test.docx")
        sc = StrategyCard()
        assert _render_draft_documents(sc, out, ctx) is True
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "答辩状" in text

    def test_plaintiff_generates_complaint(self, tmp_path):
        ctx = PipelineContext()
        ctx.identity = "起诉方"
        ctx.goal = "起诉立案"
        ctx.fact_card = FactCard(
            case_id="C1",
            parties=[Party(name="原告名", role="原告"), Party(name="被告名", role="被告")],
            key_facts=["事实1"],
            amount="10万元",
        )
        out = str(tmp_path / "test.docx")
        sc = StrategyCard()
        assert _render_draft_documents(sc, out, ctx) is True

    def test_company_defendant_format(self, tmp_path):
        ctx = PipelineContext()
        ctx.identity = "被诉方"
        ctx.fact_card = FactCard(
            parties=[Party(name="上海嘉忞贸易有限公司", role="被告")],
        )
        out = str(tmp_path / "test.docx")
        sc = StrategyCard()
        _render_draft_documents(sc, out, ctx)
        from docx import Document
        text = "\n".join(p.text for p in Document(out).paragraphs)
        assert "统一社会信用代码" in text or "住所地" in text

    def test_with_draft_documents(self, tmp_path):
        ctx = PipelineContext()
        ctx.identity = "被诉方"
        ctx.fact_card = FactCard(parties=[Party(name="张三", role="被告")])
        out = str(tmp_path / "test.docx")
        sc = StrategyCard(draft_documents=[
            DraftDocument(doc_type="答辩状", title="答辩状", content="答辩内容"),
        ])
        assert _render_draft_documents(sc, out, ctx) is True

    def test_generic_identity(self, tmp_path):
        ctx = PipelineContext()
        ctx.identity = "整理证据"
        ctx.fact_card = FactCard()
        out = str(tmp_path / "test.docx")
        sc = StrategyCard(
            situation_assessment="评估",
            action_advice=[ActionAdvice(action="做X", priority="S")],
            evidence_gap=["缺口1"],
            risk_warnings=["风险1"],
        )
        assert _render_draft_documents(sc, out, ctx) is True


# ══════════════════════════════════════════════════════════════════════
# _manual_build_zip
# ══════════════════════════════════════════════════════════════════════
class TestManualBuildZip:
    def test_empty_dir(self, tmp_path):
        src = tmp_path / "src"
        src.mkdir()
        out = str(tmp_path / "out.zip")
        _manual_build_zip(str(src), out)
        assert os.path.exists(out)

    def test_with_files(self, tmp_path):
        src = tmp_path / "src"
        src.mkdir()
        (src / "a.txt").write_text("test", encoding="utf-8")
        (src / "b.docx").write_bytes(b"fake")
        out = str(tmp_path / "out.zip")
        _manual_build_zip(str(src), out)
        import zipfile
        with zipfile.ZipFile(out) as zf:
            assert len(zf.namelist()) == 2

    def test_nested_dirs(self, tmp_path):
        src = tmp_path / "src"
        sub = src / "sub"
        sub.mkdir(parents=True)
        (sub / "deep.txt").write_text("deep", encoding="utf-8")
        out = str(tmp_path / "out.zip")
        _manual_build_zip(str(src), out)
        import zipfile
        with zipfile.ZipFile(out) as zf:
            names = zf.namelist()
            assert any("deep" in n for n in names)


# ══════════════════════════════════════════════════════════════════════
# _save_distilled_card_as_text
# ══════════════════════════════════════════════════════════════════════
class TestSaveDistilledCardAsText:
    def test_full_data(self, tmp_path):
        ctx = PipelineContext()
        dc = DistilledCard(
            fact_card=FactCard(
                case_id="C1", court="法院", identity="被告",
                amount="10万", parties=[Party(name="张三", role="被告")],
                key_facts=["事实1"], disputed_facts=["争议1"],
                conflicts=["冲突1"], missing_materials=["材料1"],
            ),
            strategy_card=StrategyCard(
                sabcd_rating="B", situation_assessment="评估",
                action_advice=[ActionAdvice(action="做X", priority="S", reasoning="因为Y")],
                evidence_gap=["缺口1"], risk_warnings=["风险1"],
                draft_documents=[DraftDocument(doc_type="答辩状", title="答辩状", content="内容")],
            ),
        )
        out = str(tmp_path / "report.txt")
        _save_distilled_card_as_text(dc, out, ctx)
        text = open(out, encoding="utf-8").read()
        assert "C1" in text
        assert "B" in text
        assert "张三" in text

    def test_empty_cards(self, tmp_path):
        ctx = PipelineContext()
        dc = DistilledCard(fact_card=FactCard(), strategy_card=StrategyCard())
        out = str(tmp_path / "report.txt")
        _save_distilled_card_as_text(dc, out, ctx)
        assert os.path.exists(out)

    def test_creates_dir(self, tmp_path):
        ctx = PipelineContext()
        dc = DistilledCard(fact_card=FactCard(case_id="C1"))
        out = str(tmp_path / "new_dir" / "report.txt")
        _save_distilled_card_as_text(dc, out, ctx)
        assert os.path.exists(out)


# ══════════════════════════════════════════════════════════════════════
# is_retryable
# ══════════════════════════════════════════════════════════════════════
class TestIsRetryable:
    def test_retryable_codes(self):
        for code in ["PermissionError", "WinError", "OSError", "TimeoutError",
                      "ConnectionError", "RENDER_FAILED", "EMPTY_FILE", "ZIP_WRITE_ERROR"]:
            assert is_retryable(code) is True, f"{code} should be retryable"

    def test_non_retryable_codes(self):
        for code in ["FileNotFoundError", "ModuleNotFoundError", "ImportError",
                      "ENOSPC", "TEMPLATE_MISSING", "SOURCE_MISSING", "ENVIRONMENT_ERROR"]:
            assert is_retryable(code) is False, f"{code} should NOT be retryable"

    def test_msg_no_space_left(self):
        assert is_retryable("OSError", "No space left on device") is False

    def test_msg_enospc(self):
        assert is_retryable("OSError", "ENOSPC") is False

    def test_msg_no_such_file(self):
        assert is_retryable("OSError", "No such file or directory") is False

    def test_msg_not_found_cn(self):
        assert is_retryable("OSError", "找不到文件") is False

    def test_msg_not_installed(self):
        assert is_retryable("Error", "reportlab not installed") is False

    def test_msg_not_found(self):
        assert is_retryable("Error", "module not found") is False

    def test_unknown_code_default_retryable(self):
        assert is_retryable("SomeNewError") is True

    def test_empty_args(self):
        assert is_retryable("") is True


# ══════════════════════════════════════════════════════════════════════
# _file_size
# ══════════════════════════════════════════════════════════════════════
class TestFileSize:
    def test_existing_file(self, tmp_path):
        p = tmp_path / "test.txt"
        p.write_bytes(b"hello")
        assert _file_size(str(p)) == 5

    def test_nonexistent(self):
        assert _file_size("/nonexistent/path") == 0

    def test_empty_file(self, tmp_path):
        p = tmp_path / "empty.txt"
        p.write_bytes(b"")
        assert _file_size(str(p)) == 0


# ══════════════════════════════════════════════════════════════════════
# _check_source_dependency
# ══════════════════════════════════════════════════════════════════════
class TestCheckSourceDependency:
    def test_no_source(self, tmp_path):
        ok, reason = _check_source_dependency({}, str(tmp_path))
        assert ok is True
        assert reason == ""

    def test_empty_source(self, tmp_path):
        ok, reason = _check_source_dependency({"source_file": ""}, str(tmp_path))
        assert ok is True

    def test_source_exists(self, tmp_path):
        src = tmp_path / "source.docx"
        src.write_bytes(b"fake")
        ok, reason = _check_source_dependency({"source_file": str(src)}, str(tmp_path))
        assert ok is True

    def test_source_missing(self, tmp_path):
        ok, reason = _check_source_dependency(
            {"source_file": str(tmp_path / "missing.docx")}, str(tmp_path)
        )
        assert ok is False
        assert "不存在" in reason

    def test_source_empty_file(self, tmp_path):
        src = tmp_path / "empty.docx"
        src.write_bytes(b"")
        ok, reason = _check_source_dependency({"source_file": str(src)}, str(tmp_path))
        assert ok is False
        assert "为空" in reason


# ══════════════════════════════════════════════════════════════════════
# render_with_manifest
# ══════════════════════════════════════════════════════════════════════
class TestRenderWithManifest:
    def _make_ts(self):
        """Create a mock-like TaskStore backed by real SQLite."""
        from core.task_store import get_task_store
        return get_task_store()

    def test_render_success(self, tmp_path):
        ctx = PipelineContext()
        ts = self._make_ts()
        task_id = "test_manifest_001"
        out = str(tmp_path / "test.txt")

        def render_fn():
            with open(out, "w") as f:
                f.write("content")
            return True

        result = render_with_manifest(ctx, task_id, "test.txt", "txt", render_fn, out, ts=ts)
        assert result is True
        assert os.path.exists(out)

    def test_render_failure(self, tmp_path):
        ctx = PipelineContext()
        ts = self._make_ts()
        task_id = "test_manifest_002"
        out = str(tmp_path / "test.txt")

        def render_fn():
            return False

        result = render_with_manifest(ctx, task_id, "test.txt", "txt", render_fn, out, ts=ts)
        assert result is False

    def test_render_exception(self, tmp_path):
        ctx = PipelineContext()
        ts = self._make_ts()
        task_id = "test_manifest_003"
        out = str(tmp_path / "test.txt")

        def render_fn():
            raise ValueError("test error")

        result = render_with_manifest(ctx, task_id, "test.txt", "txt", render_fn, out, ts=ts)
        assert result is False

    def test_render_empty_file(self, tmp_path):
        ctx = PipelineContext()
        ts = self._make_ts()
        task_id = "test_manifest_004"
        out = str(tmp_path / "test.txt")

        def render_fn():
            with open(out, "w") as f:
                pass  # empty
            return True

        result = render_with_manifest(ctx, task_id, "test.txt", "txt", render_fn, out, ts=ts)
        assert result is False  # empty file = failure

    def test_skip_already_success(self, tmp_path):
        ctx = PipelineContext()
        ts = self._make_ts()
        task_id = "test_manifest_005"
        out = tmp_path / "test.txt"
        out.write_text("existing", encoding="utf-8")

        # Pre-mark as success
        ts.manifest_init_entry(task_id, "test.txt", "txt")
        ts.manifest_mark_success(task_id, "test.txt", file_size=8)

        call_count = [0]
        def render_fn():
            call_count[0] += 1
            return True

        result = render_with_manifest(ctx, task_id, "test.txt", "txt", render_fn, str(out), ts=ts)
        assert result is True
        assert call_count[0] == 0  # render_fn was NOT called
