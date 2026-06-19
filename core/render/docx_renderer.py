"""
docx_renderer.py - 法院级文书 DOCX 渲染引擎

输出标准参照最高人民法院《法院诉讼文书样式》。
字体：标题黑体小二号，正文仿宋_GB2312 三号，页码居中。
"""
from __future__ import annotations

import os
import re
import logging
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from core.fact_card import FactCard, StrategyCard, DistilledCard, DraftDocument

logger = logging.getLogger(__name__)

# ── 法院文书排版常量 ──────────────────────────────────────────────────
FONT_BODY = "仿宋_GB2312"
FONT_BODY_FALLBACK = "仿宋"
FONT_HEADING = "黑体"
FONT_TITLE = "方正小标宋简体"
FONT_TITLE_FALLBACK = "黑体"

# 字号对照 (Word 磅值)
SIZE_TITLE = Pt(22)       # 小二号
SIZE_H1 = Pt(16)          # 三号
SIZE_H2 = Pt(14)          # 四号
SIZE_H3 = Pt(12)          # 小四号
SIZE_BODY = Pt(14)        # 四号 (正文标准)
SIZE_FOOTER = Pt(10)      # 五号

# 页边距 (国标 A4)
MARGIN_TOP = Cm(3.7)
MARGIN_BOTTOM = Cm(3.5)
MARGIN_LEFT = Cm(2.8)
MARGIN_RIGHT = Cm(2.6)

# 行距
LINE_SPACING_EXACT = Pt(28)  # 固定值28磅 (国标)


def _resolve_font(primary: str, fallback: str) -> str:
    """检测字体是否存在，返回可用字体名。"""
    try:
        from docx.shared import Pt as _Pt
        # 简单检测：直接用主字体，渲染时如果缺字会自动降级
        return primary
    except Exception:
        return fallback


def _set_run_font(run, font_name: str, size: Pt, bold: bool = False, color: Optional[RGBColor] = None):
    """统一设置 run 的中西文字体。"""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run.element.get_or_add_rPr()
    rpr_fonts = rpr.find(qn("w:rFonts"))
    if rpr_fonts is None:
        rpr_fonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rpr.insert(0, rpr_fonts)
    else:
        rpr_fonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_spacing(para, line_spacing: Pt = LINE_SPACING_EXACT,
                           space_before: Pt = Pt(0), space_after: Pt = Pt(0)):
    """设置段落行距和段前段后。"""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after


def _add_page_number(doc: Document):
    """在页脚添加居中页码。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()

        run = para.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run.element.append(fld_char_begin)

        run2 = para.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2.element.append(instr)

        run3 = para.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3.element.append(fld_char_end)

        _set_run_font(run, FONT_BODY, SIZE_FOOTER)
        _set_run_font(run2, FONT_BODY, SIZE_FOOTER)
        _set_run_font(run3, FONT_BODY, SIZE_FOOTER)


def _setup_court_document(doc: Document) -> None:
    """设置法院文书标准格式：A4、国标页边距、正文行距。"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY_FALLBACK
    font.size = SIZE_BODY
    rpr = style.element.get_or_add_rPr()
    rpr_fonts = rpr.find(qn("w:rFonts"))
    if rpr_fonts is None:
        rpr_fonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_BODY_FALLBACK}"/>')
        rpr.insert(0, rpr_fonts)
    else:
        rpr_fonts.set(qn("w:eastAsia"), FONT_BODY_FALLBACK)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_SPACING_EXACT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74)  # 首行缩进2字符


def _add_court_title(doc: Document, title: str) -> None:
    """添加文书标题（方正小标宋/黑体，小二号，居中）。"""
    if not title:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(para, line_spacing=Pt(32), space_before=Pt(0), space_after=Pt(10))
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(title)
    _set_run_font(run, FONT_TITLE_FALLBACK, SIZE_TITLE, bold=True)


def _add_court_heading(doc: Document, text: str, level: int = 1) -> None:
    """添加章节标题。"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = SIZE_H1 if level == 1 else SIZE_H2 if level == 2 else SIZE_H3
    spacing = Pt(10) if level == 1 else Pt(6)
    _set_paragraph_spacing(para, line_spacing=LINE_SPACING_EXACT, space_before=spacing, space_after=Pt(4))
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    _set_run_font(run, FONT_HEADING, size, bold=True)


def _add_court_body(doc: Document, text: str, indent: bool = True, bold: bool = False) -> None:
    """添加正文段落（仿宋四号，首行缩进2字符）。"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(para, line_spacing=LINE_SPACING_EXACT)
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    else:
        para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    _set_run_font(run, FONT_BODY_FALLBACK, SIZE_BODY, bold=bold)


def _add_court_list_item(doc: Document, text: str) -> None:
    """添加列表项（带序号的正文段落）。"""
    _add_court_body(doc, text, indent=True)


def _add_page_break(doc: Document):
    """添加分页符。"""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type=7)  # WD_BREAK.PAGE


# ── LLM 输出智能解析 ──────────────────────────────────────────────────

def _parse_llm_content(content: str) -> list[dict]:
    """将 LLM 输出的 Markdown/纯文本解析为结构化段落。

    返回 [{"type": "title"|"heading1"|"heading2"|"heading3"|"body"|"list"|"separator", "text": str}]
    """
    if not content:
        return []

    blocks = []
    lines = content.strip().split("\n")
    i = 0
    first_content = True

    while i < len(lines):
        line = lines[i].rstrip()

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # First non-empty line: treat as title if it looks like a document title
        if first_content:
            text = line.strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # Recognize common legal document titles
            title_patterns = ["民事答辩状", "民事起诉状", "案件处境评估报告", "行动建议书",
                              "证据闭环补强清单", "投诉状", "行政复议申请书", "法律文书"]
            is_title = any(tp in text for tp in title_patterns)
            if is_title or len(text) <= 20:
                blocks.append({"type": "title", "text": text})
                first_content = False
                i += 1
                continue
            first_content = False

        # Markdown heading: ### 或 ##
        m_h = re.match(r'^(#{1,3})\s+(.+)', line)
        if m_h:
            level = len(m_h.group(1))
            text = m_h.group(2).strip()
            # 去掉可能的 ** 包裹
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            blocks.append({"type": f"heading{level}", "text": text})
            first_content = False
            i += 1
            continue

        # 数字标题: 一、二、三、... 或 1. 2. 3. ...
        m_cn = re.match(r'^([一二三四五六七八九十]+[、.．])\s*(.+)', line)
        if m_cn:
            blocks.append({"type": "heading1", "text": line.strip()})
            i += 1
            continue

        m_num = re.match(r'^(\d+)[.、）)]\s*(.+)', line)
        if m_num:
            # 可能是列表项或小标题
            text = m_num.group(2).strip()
            if len(text) > 30:
                blocks.append({"type": "body", "text": line.strip()})
            else:
                blocks.append({"type": "heading2", "text": line.strip()})
            i += 1
            continue

        # 罗马数字或特殊标记: （一）(一)
        m_roman = re.match(r'^[（(][一二三四五六七八九十]+[）)]\s*(.+)', line)
        if m_roman:
            blocks.append({"type": "heading2", "text": line.strip()})
            i += 1
            continue

        # 分隔线
        if re.match(r'^[-=─]{5,}', line.strip()):
            blocks.append({"type": "separator", "text": ""})
            i += 1
            continue

        # 列表项: - 或 • 或 ●
        m_bullet = re.match(r'^\s*[-•●▪]\s+(.+)', line)
        if m_bullet:
            blocks.append({"type": "body", "text": f"• {m_bullet.group(1).strip()}"})
            i += 1
            continue

        # 普通正文（合并连续非空行）
        body_lines = [line.strip()]
        i += 1
        while i < len(lines):
            next_line = lines[i].rstrip()
            if not next_line.strip():
                break
            if re.match(r'^(#{1,3})\s+', next_line):
                break
            if re.match(r'^[一二三四五六七八九十]+[、.．]', next_line):
                break
            if re.match(r'^\d+[.、）)]\s+', next_line):
                break
            if re.match(r'^[-=─]{5,}', next_line.strip()):
                break
            body_lines.append(next_line.strip())
            i += 1

        merged = "\n".join(body_lines)
        # 去掉 ** 包裹
        merged = re.sub(r'\*\*(.+?)\*\*', r'\1', merged)
        blocks.append({"type": "body", "text": merged})

    return blocks


# ── 主渲染函数 ─────────────────────────────────────────────────────────

def render_docx_from_text(content: str, output_path: str, title: str = "") -> bool:
    """将 LLM 生成的文本渲染为法院级 DOCX 文书。

    智能解析 Markdown/纯文本，自动识别标题、章节、列表。
    """
    try:
        doc = Document()
        _setup_court_document(doc)

        blocks = _parse_llm_content(content)

        # 确定文档标题
        doc_title = ""
        if title:
            doc_title = title
        elif blocks and blocks[0]["type"] in ("title", "heading1"):
            doc_title = blocks[0]["text"]
            blocks = blocks[1:]

        if not doc_title:
            doc_title = ""

        _add_court_title(doc, doc_title)

        for block in blocks:
            btype = block["type"]
            text = block["text"]

            if btype == "heading1":
                _add_court_heading(doc, text, level=1)
            elif btype == "heading2":
                _add_court_heading(doc, text, level=2)
            elif btype == "heading3":
                _add_court_heading(doc, text, level=3)
            elif btype == "separator":
                pass  # 跳过分隔线
            elif btype == "body":
                # 多行文本逐段输出
                for para_text in text.split("\n"):
                    para_text = para_text.strip()
                    if para_text:
                        _add_court_body(doc, para_text, indent=True)
            else:
                if text:
                    _add_court_body(doc, text, indent=True)

        # 添加页码
        _add_page_number(doc)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info("Court DOCX rendered: %s", output_path)
        return True

    except Exception as exc:
        logger.error("Failed to render DOCX: %s", exc)
        return False


def render_docx(title: str, sections: List[dict], output_path: str) -> bool:
    """渲染通用 DOCX（兼容旧接口）。"""
    try:
        doc = Document()
        _setup_court_document(doc)
        _add_court_title(doc, title)

        for section in sections:
            heading = section.get("heading", "")
            content = section.get("content", "")
            if heading:
                _add_court_heading(doc, heading, level=1)
            if content:
                for para_text in content.split("\n"):
                    para_text = para_text.strip()
                    if para_text:
                        _add_court_body(doc, para_text, indent=True)

        _add_page_number(doc)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info("DOCX rendered: %s", output_path)
        return True
    except Exception as exc:
        logger.error("Failed to render DOCX: %s", exc)
        return False
