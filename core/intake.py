"""
intake.py - File Discovery and Reading

Scans the input directory for legal documents and extracts text content.
Supports: PDF, DOCX, DOC, TXT, JPG, PNG
Uses pypdf for PDF, python-docx for DOCX, plain open() for TXT.
Images are OCR'd via MiMo 2.5 multimodal vision API (priority) or DeepSeek.
"""
from __future__ import annotations
import logging
import os
import glob as globmod
from pathlib import Path
from typing import List

from core.fact_card import PipelineContext

logger = logging.getLogger(__name__)


# Supported file extensions and their types
SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc",
    ".txt": "txt",
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
}


def scan_input_dir(input_dir: str) -> List[str]:
    """
    Find all supported files in the input directory.
    Searches recursively through subdirectories.
    Returns list of absolute file paths sorted by name.
    """
    if not os.path.isdir(input_dir):
        return []

    found_files: List[str] = []
    for ext in SUPPORTED_EXTENSIONS:
        pattern = os.path.join(input_dir, "**", f"*{ext}")
        matches = globmod.glob(pattern, recursive=True)
        found_files.extend(matches)

    # Also search for uppercase extensions
    for ext in SUPPORTED_EXTENSIONS:
        pattern = os.path.join(input_dir, "**", f"*{ext.upper()}")
        matches = globmod.glob(pattern, recursive=True)
        found_files.extend(matches)

    # Deduplicate (in case both .pdf and .PDF match) and sort
    unique_files = sorted(set(os.path.abspath(f) for f in found_files))
    return unique_files


def read_text_file(path: str) -> str:
    """Read a .txt file and return its content as a string."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def read_docx_text(path: str) -> str:
    """
    Extract all text from a .docx file using python-docx.
    Falls back to basic XML extraction if python-docx fails.
    """
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n".join(paragraphs)
    except Exception:
        # Fallback: try to read raw XML from the docx zip
        import zipfile
        import xml.etree.ElementTree as ET

        try:
            with zipfile.ZipFile(path, "r") as z:
                with z.open("word/document.xml") as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    texts = []
                    for t in root.iter(f"{{{ns['w']}}}t"):
                        if t.text:
                            texts.append(t.text)
                    return " ".join(texts)
        except Exception:
            return f"[无法读取DOCX文件: {os.path.basename(path)}]"


def read_pdf_text(path: str) -> str:
    """
    Extract text from a PDF file using pypdf.
    Returns extracted text with page separators.
    Returns empty string for scanned PDFs (no extractable text).
    """
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(f"--- 第{i + 1}页 ---\n{text.strip()}")
        result = "\n\n".join(pages_text)
        if not result.strip():
            return ""  # Scanned PDF — no extractable text, needs OCR
        return result
    except ImportError:
        return f"[无法读取PDF文件: pypdf未安装 - {os.path.basename(path)}]"
    except Exception as e:
        return f"[PDF读取失败: {os.path.basename(path)} - {str(e)}]"


def extract_pdf_page_images(pdf_path: str, output_dir: str = None) -> list[str]:
    """Extract each page of a PDF as a PNG image for OCR.

    Uses pdf2image (poppler) or falls back to PyMuPDF (fitz).
    Returns list of image file paths. Max 20 pages to avoid timeout.
    """
    import tempfile

    MAX_PAGES = 20
    DPI = 150  # Lower DPI for faster processing

    if output_dir is None:
        output_dir = tempfile.mkdtemp(prefix="pdf_ocr_")
    os.makedirs(output_dir, exist_ok=True)

    image_paths = []

    # Try pdf2image first (requires poppler)
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(pdf_path, dpi=DPI, fmt="png", last_page=MAX_PAGES)
        for i, img in enumerate(images):
            img_path = os.path.join(output_dir, f"page_{i+1}.png")
            img.save(img_path, "PNG")
            image_paths.append(img_path)
        logger.info("PDF→图片: pdf2image, %d pages, DPI=%d", len(image_paths), DPI)
        return image_paths
    except (ImportError, Exception):
        pass

    # Try PyMuPDF (fitz)
    try:
        import fitz
        doc = fitz.open(pdf_path)
        total = min(len(doc), MAX_PAGES)
        for i in range(total):
            page = doc[i]
            pix = page.get_pixmap(dpi=DPI)
            img_path = os.path.join(output_dir, f"page_{i+1}.png")
            pix.save(img_path)
            image_paths.append(img_path)
        doc.close()
        logger.info("PDF→图片: PyMuPDF, %d pages, DPI=%d", len(image_paths), DPI)
        return image_paths
    except (ImportError, Exception):
        pass

    # Try subprocess with magick/convert
    try:
        import subprocess
        for i in range(MAX_PAGES):
            img_path = os.path.join(output_dir, f"page_{i+1}.png")
            result = subprocess.run(
                ["magick", "-density", str(DPI), f"{pdf_path}[{i}]", img_path],
                capture_output=True, timeout=30,
            )
            if result.returncode != 0 or not os.path.exists(img_path):
                break
            image_paths.append(img_path)
        logger.info("PDF→图片: magick, %d pages, DPI=%d", len(image_paths), DPI)
        return image_paths
    except Exception:
        pass

    return []


def read_image_placeholder(path: str) -> str:
    """
    Return a placeholder for image files.
    OCR is not implemented at this stage - images are noted for later processing.
    """
    filename = os.path.basename(path)
    return f"[图片文件: {filename} - 需要OCR处理才能提取文字内容]"


def ocr_images_via_api(image_paths: list[str]) -> dict[str, str]:
    """Use vision API to OCR image files and return {path: text} mapping.

    Priority: MiMo 2.5 (multimodal vision) > DeepSeek (text-only fallback).
    MiMo 2.5 supports native image understanding and is preferred for OCR.
    Includes retry logic (max 3 attempts per batch) for API transient failures.
    """
    import time

    MAX_RETRIES = 3
    RETRY_DELAY = 2  # seconds

    try:
        from core.settings_store import get_settings_store
        store = get_settings_store()
        s = store.settings

        provider = None
        vision_model = ""

        # MiMo 2.5 has native multimodal/vision support — highest priority
        if s.mimo.is_configured():
            from core.ai.unified_client import UnifiedAIClient, ProviderConfig
            provider = UnifiedAIClient(ProviderConfig(
                name="mimo",
                api_key=s.mimo.api_key,
                base_url=s.mimo.base_url,
                model=s.mimo.model,
                timeout=max(s.mimo.timeout, 120),
            ))
            vision_model = s.mimo.model
            logger.info("OCR: using MiMo vision model=%s", vision_model)
        elif s.deepseek.is_configured():
            from core.ai.unified_client import UnifiedAIClient, ProviderConfig
            provider = UnifiedAIClient(ProviderConfig(
                name="deepseek",
                api_key=s.deepseek.api_key,
                base_url=s.deepseek.base_url,
                model="deepseek-chat",
                timeout=max(s.deepseek.timeout, 120),
            ))
            vision_model = "deepseek-chat"
            logger.info("OCR: using DeepSeek fallback model=%s", vision_model)

        if not provider or not provider.is_configured:
            logger.warning("OCR: no API provider configured, skipping vision OCR")
            return {}

        from core.ai.multimodal import encode_image_base64
        results = {}
        batch_size = 3
        total_batches = (len(image_paths) + batch_size - 1) // batch_size

        for i in range(0, len(image_paths), batch_size):
            batch = image_paths[i:i + batch_size]
            batch_num = i // batch_size + 1
            content = [{"type": "text", "text": "请逐张分析这些证据图片，提取所有可见文字内容、当事人信息、日期、金额等关键法律信息。对每张图片分别输出，标注文件名。"}]

            for img_path in batch:
                b64 = encode_image_base64(img_path)
                if b64:
                    content.append({"type": "image_url", "image_url": {"url": b64}})
                else:
                    logger.warning("OCR: failed to encode image %s", img_path)

            if len(content) <= 1:
                logger.warning("OCR: batch %d/%d has no valid images, skipping", batch_num, total_batches)
                continue

            messages = [
                {"role": "system", "content": "你是一个法律文档OCR助手。请精确提取图片中的所有文字内容。"},
                {"role": "user", "content": content},
            ]

            # Retry logic
            response = None
            for attempt in range(1, MAX_RETRIES + 1):
                response = provider.chat_messages(messages, model=vision_model, temperature=0.1, max_tokens=4096)
                if response.success and response.content:
                    break
                if attempt < MAX_RETRIES:
                    err_msg = response.error if response else "unknown"
                    logger.warning("OCR: batch %d/%d attempt %d failed: %s, retrying in %ds",
                                   batch_num, total_batches, attempt, err_msg, RETRY_DELAY)
                    time.sleep(RETRY_DELAY)
                    RETRY_DELAY *= 2  # Exponential backoff

            if response.success and response.content:
                parts = response.content.split("---")
                for j, img_path in enumerate(batch):
                    if j < len(parts):
                        results[img_path] = parts[j].strip()
                    else:
                        results[img_path] = response.content.strip()
                logger.info("OCR: batch %d/%d succeeded, extracted %d images",
                            batch_num, total_batches, len(batch))
            else:
                err_msg = response.error if response else "unknown"
                logger.error("OCR: batch %d/%d failed after %d attempts (model=%s): %s",
                             batch_num, total_batches, MAX_RETRIES, vision_model, err_msg)

        logger.info("OCR: total %d/%d images extracted successfully",
                     len(results), len(image_paths))
        return results
    except Exception as exc:
        logger.error("OCR: vision API call failed with exception: %s", exc, exc_info=True)
        return {}


def read_file(path: str) -> str:
    """
    Read a file based on its extension.
    Returns the extracted text content or an error message.
    """
    ext = os.path.splitext(path)[1].lower()
    file_type = SUPPORTED_EXTENSIONS.get(ext, "unknown")

    try:
        if file_type == "txt":
            return read_text_file(path)
        elif file_type == "docx":
            return read_docx_text(path)
        elif file_type == "pdf":
            return read_pdf_text(path)
        elif file_type == "image":
            return read_image_placeholder(path)
        elif file_type == "doc":
            # Legacy .doc format - attempt with python-docx (may fail)
            try:
                return read_docx_text(path)
            except Exception:
                return f"[无法读取DOC文件: {os.path.basename(path)} - 请转换为DOCX格式]"
        else:
            return f"[不支持的文件格式: {ext} - {os.path.basename(path)}]"
    except Exception as e:
        return f"[文件读取错误: {os.path.basename(path)} - {str(e)}]"


def run_intake(ctx: PipelineContext) -> PipelineContext:
    """
    Main entry point for the intake pipeline stage.
    
    1. Scans input_dir for supported files
    2. Reads text content from each file
    3. Populates ctx.raw_texts and ctx.file_list
    4. Logs each file found and processed
    5. Skips unreadable files with warning
    """
    ctx.log(f"开始扫描输入目录: {ctx.input_dir}")

    if not os.path.isdir(ctx.input_dir):
        ctx.add_error(f"输入目录不存在: {ctx.input_dir}")
        return ctx

    files = scan_input_dir(ctx.input_dir)

    if not files:
        ctx.add_error(f"输入目录中未找到支持的文件: {ctx.input_dir}")
        return ctx

    ctx.log(f"找到 {len(files)} 个文件")

    ctx.file_list = files
    ctx.raw_texts = []

    image_files = [f for f in files
                   if os.path.splitext(f)[1].lower() in {'.jpg', '.jpeg', '.png'}]
    text_files = [f for f in files if f not in image_files]

    for file_path in text_files:
        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()
        file_type = SUPPORTED_EXTENSIONS.get(ext, "unknown")

        ctx.log(f"读取文件: {filename} (类型: {file_type})")

        text = read_file(file_path)

        # Detect scanned PDF (pypdf returns empty text)
        if file_type == "pdf" and not text.strip():
            ctx.log(f"  {filename}: 文本提取为空，尝试PDF页面截图OCR...")
            page_images = extract_pdf_page_images(file_path)
            if page_images:
                ctx.log(f"  {filename}: 提取了 {len(page_images)} 页图片，开始逐页OCR...")
                ocr_for_pdf = ocr_images_via_api(page_images)
                combined = "\n\n".join(
                    f"--- 第{i+1}页 ---\n{ocr_for_pdf.get(p, '')}"
                    for i, p in enumerate(page_images)
                    if ocr_for_pdf.get(p, "").strip()
                )
                if combined.strip():
                    text = combined
                    ctx.log(f"  {filename}: PDF-OCR成功 ({len(text)} 字符, {len(ocr_for_pdf)}页)")
                else:
                    ctx.log(f"  {filename}: PDF-OCR未提取到文字")
                # Cleanup temp images
                for p in page_images:
                    try:
                        os.remove(p)
                    except Exception:
                        pass
                try:
                    os.rmdir(os.path.dirname(page_images[0]))
                except Exception:
                    pass

        if text.startswith("[") and text.endswith("]"):
            ctx.log(f"警告: {filename} - {text}")
            ctx.raw_texts.append(text)
        else:
            ctx.raw_texts.append(text)
            word_count = len(text)
            ctx.log(f"成功读取: {filename} ({word_count} 字符)")

    if image_files:
        ctx.log(f"检测到 {len(image_files)} 个图片文件，正在通过 Vision API 提取文字...")
        ocr_results = ocr_images_via_api(image_files)

        for img_path in image_files:
            filename = os.path.basename(img_path)
            if img_path in ocr_results and ocr_results[img_path].strip():
                text = ocr_results[img_path]
                ctx.raw_texts.append(text)
                ctx.log(f"OCR 成功: {filename} ({len(text)} 字符)")
            else:
                placeholder = read_image_placeholder(img_path)
                ctx.raw_texts.append(placeholder)
                ctx.log(f"OCR 失败/跳过: {filename}，使用占位符")

    ctx.log(f"文件读取完成: 成功 {sum(1 for t in ctx.raw_texts if not t.startswith('['))} 个, "
            f"失败/跳过 {sum(1 for t in ctx.raw_texts if t.startswith('['))} 个")

    # Guard: if ALL raw_texts are placeholders (e.g. all images with failed OCR),
    # the pipeline will produce empty results downstream. Surface the error now.
    real_texts = [t for t in ctx.raw_texts if not t.startswith("[")]
    if not real_texts and image_files:
        ctx.add_error(
            f"所有 {len(image_files)} 张图片均未能提取文字内容。"
            f"请检查：1) AI 服务是否已配置且支持图片识别；"
            f"2) 图片是否清晰可读。"
            f"建议：将判决书转换为 PDF 格式后重新上传。"
        )

    return ctx
