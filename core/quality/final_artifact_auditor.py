"""
final_artifact_auditor.py - Quality Gate: Final Artifact Audit

The FINAL arbiter of customer output quality. Runs 8 comprehensive checks
on the customer output directory to ensure all artifacts meet quality standards.

If this audit fails, the whole pipeline fails.
"""
from __future__ import annotations

import os
from dataclasses import dataclass, field
from typing import List

# NOTE: Imports from visible_docx_checker and package_leak_scanner are done
# lazily inside functions to avoid circular imports (they import CheckResult
# from this module).


@dataclass
class CheckResult:
    """Result of a single quality check."""
    check_name: str = ""
    passed: bool = False
    message: str = ""
    details: str = ""


@dataclass
class AuditReport:
    """Aggregated results of all quality checks on a customer directory."""
    case_dir: str = ""
    checks: List[CheckResult] = field(default_factory=list)

    @property
    def passed(self) -> bool:
        return all(c.passed for c in self.checks)

    @property
    def failed_checks(self) -> List[CheckResult]:
        return [c for c in self.checks if not c.passed]

    def summary(self) -> str:
        total = len(self.checks)
        passed = sum(1 for c in self.checks if c.passed)
        return f"Quality Gate: {passed}/{total} checks passed"


# ---------------------------------------------------------------------------
# V18-RC expected document names in customer output
# ---------------------------------------------------------------------------
EXPECTED_DOCX_NAMES = [
    "01_案件处境评估报告.docx",
    "02_行动建议书.docx",
    "03_证据闭环补强清单.docx",
    "05_可提交文书草稿.docx",
    "06_答辩状.docx",
]

# Legacy names (for backward compatibility check)
LEGACY_DOCX_NAMES = [
    "事实与证据清单.docx",
    "法律分析报告.docx",
    "策略建议书.docx",
    "起诉状.docx",
    "证据目录.docx",
]


def _count_chinese(text: str) -> int:
    """Count CJK characters in text."""
    return sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")


def _read_docx_text(docx_path: str) -> str:
    """Extract all text from a DOCX file. Returns empty string on failure."""
    try:
        from docx import Document

        doc = Document(docx_path)
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)
        return "\n".join(parts)
    except Exception:
        return ""


def _find_docx_files(customer_dir: str) -> List[str]:
    """Find all .docx files under customer_dir (non-recursive)."""
    results: List[str] = []
    if not os.path.isdir(customer_dir):
        return results
    for name in os.listdir(customer_dir):
        if name.lower().endswith(".docx") and os.path.isfile(
            os.path.join(customer_dir, name)
        ):
            results.append(os.path.join(customer_dir, name))
    return results


def _find_pdf_files(customer_dir: str) -> List[str]:
    """Find all .pdf files under customer_dir (non-recursive)."""
    results: List[str] = []
    if not os.path.isdir(customer_dir):
        return results
    for name in os.listdir(customer_dir):
        if name.lower().endswith(".pdf") and os.path.isfile(
            os.path.join(customer_dir, name)
        ):
            results.append(os.path.join(customer_dir, name))
    return results


def _find_xlsx_files(customer_dir: str) -> List[str]:
    """Find all .xlsx files under customer_dir (non-recursive)."""
    results: List[str] = []
    if not os.path.isdir(customer_dir):
        return results
    for name in os.listdir(customer_dir):
        if name.lower().endswith(".xlsx") and os.path.isfile(
            os.path.join(customer_dir, name)
        ):
            results.append(os.path.join(customer_dir, name))
    return results


# ---------------------------------------------------------------------------
# Check 1: DOCX files can be opened
# ---------------------------------------------------------------------------
def _check_docx_openable(customer_dir: str) -> CheckResult:
    """Verify every DOCX in customer_dir can be opened by python-docx."""
    from core.quality.visible_docx_checker import check_docx_readable

    docx_files = _find_docx_files(customer_dir)
    if not docx_files:
        return CheckResult(
            check_name="DOCX可打开",
            passed=False,
            message="customer目录中未发现任何DOCX文件",
        )

    failures: List[str] = []
    for fpath in docx_files:
        result = check_docx_readable(fpath)
        if not result.passed:
            failures.append(f"{os.path.basename(fpath)}: {result.message}")

    if failures:
        return CheckResult(
            check_name="DOCX可打开",
            passed=False,
            message=f"{len(failures)}个DOCX文件无法正常打开",
            details="\n".join(failures),
        )

    return CheckResult(
        check_name="DOCX可打开",
        passed=True,
        message=f"全部{len(docx_files)}个DOCX文件均可正常打开",
    )


# ---------------------------------------------------------------------------
# Check 2: DOCX has sufficient Chinese characters (>= 100 per main doc)
# ---------------------------------------------------------------------------
def _check_docx_chinese_content(customer_dir: str) -> CheckResult:
    """Verify main DOCX docs contain >= 100 Chinese characters each."""
    docx_files = _find_docx_files(customer_dir)
    if not docx_files:
        return CheckResult(
            check_name="DOCX中文内容",
            passed=False,
            message="无DOCX文件可检查",
        )

    insufficient: List[str] = []
    for fpath in docx_files:
        text = _read_docx_text(fpath)
        cn_count = _count_chinese(text)
        if cn_count < 100:
            insufficient.append(
                f"{os.path.basename(fpath)}: {cn_count}个中文字符 (不足100)"
            )

    if insufficient:
        return CheckResult(
            check_name="DOCX中文内容",
            passed=False,
            message=f"{len(insufficient)}个DOCX文件中文字符不足100个",
            details="\n".join(insufficient),
        )

    return CheckResult(
        check_name="DOCX中文内容",
        passed=True,
        message=f"全部{len(docx_files)}个DOCX文件均含>=100个中文字符",
    )


# ---------------------------------------------------------------------------
# Check 3: PDF files start with %PDF header
# ---------------------------------------------------------------------------
def _check_pdf_headers(customer_dir: str) -> CheckResult:
    """Verify every PDF starts with %PDF magic bytes."""
    from core.quality.package_leak_scanner import check_pdf_header

    pdf_files = _find_pdf_files(customer_dir)
    if not pdf_files:
        # No PDFs is not necessarily a failure - depends on pipeline config
        return CheckResult(
            check_name="PDF文件头",
            passed=True,
            message="无PDF文件需要检查",
        )

    failures: List[str] = []
    for fpath in pdf_files:
        result = check_pdf_header(fpath)
        if not result.passed:
            failures.append(f"{os.path.basename(fpath)}: {result.message}")

    if failures:
        return CheckResult(
            check_name="PDF文件头",
            passed=False,
            message=f"{len(failures)}个PDF文件头无效",
            details="\n".join(failures),
        )

    return CheckResult(
        check_name="PDF文件头",
        passed=True,
        message=f"全部{len(pdf_files)}个PDF文件头有效",
    )


# ---------------------------------------------------------------------------
# Check 4: XLSX evidence rows >= 3
# ---------------------------------------------------------------------------
def _check_xlsx_rows(customer_dir: str) -> CheckResult:
    """Verify XLSX evidence files have at least 3 data rows."""
    from core.quality.package_leak_scanner import check_xlsx_rows

    xlsx_files = _find_xlsx_files(customer_dir)
    if not xlsx_files:
        return CheckResult(
            check_name="XLSX数据行",
            passed=True,
            message="无XLSX文件需要检查",
        )

    failures: List[str] = []
    for fpath in xlsx_files:
        result = check_xlsx_rows(fpath, min_rows=3)
        if not result.passed:
            failures.append(f"{os.path.basename(fpath)}: {result.message}")

    if failures:
        return CheckResult(
            check_name="XLSX数据行",
            passed=False,
            message=f"{len(failures)}个XLSX文件数据行不足",
            details="\n".join(failures),
        )

    return CheckResult(
        check_name="XLSX数据行",
        passed=True,
        message=f"全部{len(xlsx_files)}个XLSX文件数据行>=3",
    )


# ---------------------------------------------------------------------------
# Check 5: No internal files (.json, .md, debug, ai_raw_outputs) in customer dir
# ---------------------------------------------------------------------------
def _check_no_internal_files(customer_dir: str) -> CheckResult:
    """Verify customer directory contains NO .json, .md, debug, or ai_raw_outputs files."""
    from core.quality.package_leak_scanner import scan_for_leaks

    result = scan_for_leaks(customer_dir)
    return CheckResult(
        check_name="无内部文件",
        passed=result.passed,
        message=result.message,
        details=result.details,
    )


# ---------------------------------------------------------------------------
# Check 6: No placeholder/TODO text in documents
# ---------------------------------------------------------------------------
def _check_no_placeholders(customer_dir: str) -> CheckResult:
    """Verify no DOCX contains placeholder text like TODO, 示例文本, etc."""
    from core.quality.visible_docx_checker import check_docx_no_placeholders

    docx_files = _find_docx_files(customer_dir)
    if not docx_files:
        return CheckResult(
            check_name="无占位文本",
            passed=True,
            message="无DOCX文件需要检查",
        )

    failures: List[str] = []
    for fpath in docx_files:
        result = check_docx_no_placeholders(fpath)
        if not result.passed:
            failures.append(f"{os.path.basename(fpath)}: {result.message}")

    if failures:
        return CheckResult(
            check_name="无占位文本",
            passed=False,
            message=f"{len(failures)}个DOCX含占位文本",
            details="\n".join(failures),
        )

    return CheckResult(
        check_name="无占位文本",
        passed=True,
        message=f"全部{len(docx_files)}个DOCX无占位文本",
    )


# ---------------------------------------------------------------------------
# Check 7: No internal fields in user-facing documents
# ---------------------------------------------------------------------------
def _check_no_internal_fields(customer_dir: str) -> CheckResult:
    """Verify no DOCX contains internal system fields (fact_card, prompt, etc.)."""
    from core.quality.visible_docx_checker import check_docx_no_internal_fields

    docx_files = _find_docx_files(customer_dir)
    if not docx_files:
        return CheckResult(
            check_name="无内部字段",
            passed=True,
            message="无DOCX文件需要检查",
        )

    failures: List[str] = []
    for fpath in docx_files:
        result = check_docx_no_internal_fields(fpath)
        if not result.passed:
            failures.append(f"{os.path.basename(fpath)}: {result.message}")

    if failures:
        return CheckResult(
            check_name="无内部字段",
            passed=False,
            message=f"{len(failures)}个DOCX含内部字段",
            details="\n".join(failures),
        )

    return CheckResult(
        check_name="无内部字段",
        passed=True,
        message=f"全部{len(docx_files)}个DOCX无内部字段泄露",
    )


# ---------------------------------------------------------------------------
# Check 8: All expected files exist (5 core docs + PDFs + ZIP)
# ---------------------------------------------------------------------------
def _get_expected_docx_names(identity: str = "", goal: str = "") -> List[str]:
    """Get expected DOCX names based on identity and goal."""
    base_docs = [
        "01_案件处境评估报告.docx",
        "02_行动建议书.docx",
        "03_证据闭环补强清单.docx",
        "05_可提交文书草稿.docx",
    ]
    # Goal-based extras take priority
    goal_extras = {
        "申请再审": "06_再审申请书.docx",
        "提起起诉": "06_起诉状.docx",
        "投诉举报": "06_投诉状.docx",
        "应诉答辩": "06_答辩状.docx",
        "申请行政复议": "06_行政复议申请书.docx",
        "维权投诉": "06_投诉状.docx",
        "支付令异议": "06_支付令异议书.docx",
    }
    if goal in goal_extras:
        base_docs.append(goal_extras[goal])
    else:
        identity_extras = {
            "投诉方": "06_投诉状.docx",
            "起诉方": "06_起诉状.docx",
            "起诉方（原告）": "06_起诉状.docx",
            "被诉方（被告）": "06_答辩状.docx",
            "行政复议申请人": "06_行政复议申请书.docx",
        }
        extra = identity_extras.get(identity, "")
        if extra:
            base_docs.append(extra)
    return base_docs


def _check_expected_files(customer_dir: str, identity: str = "", goal: str = "") -> CheckResult:
    """Verify all expected output files exist in customer directory."""
    missing: List[str] = []
    expected_names = _get_expected_docx_names(identity, goal)

    # Check DOCX documents - support dynamic filenames with party name
    all_files = os.listdir(customer_dir) if os.path.isdir(customer_dir) else []
    
    for doc_name in expected_names:
        # Check exact match first
        fpath = os.path.join(customer_dir, doc_name)
        if os.path.isfile(fpath):
            continue
        
        # Check for dynamic filename pattern (e.g., "05_可提交文书草稿_张三案.docx")
        base_name = doc_name.replace(".docx", "")
        found = False
        for filename in all_files:
            if not filename.endswith(".docx"):
                continue
            # Match: exact, or "base_name_*.docx" (dynamic with party name)
            if filename == doc_name or (filename.startswith(base_name + "_") or filename.startswith(base_name + "案")):
                found = True
                break
        
        if not found:
            missing.append(doc_name)

    # Check that at least one PDF exists
    pdf_files = _find_pdf_files(customer_dir)
    if not pdf_files:
        missing.append("*.pdf (至少需要一个PDF文件)")

    # Check that at least one ZIP exists
    has_zip = False
    if os.path.isdir(customer_dir):
        for name in os.listdir(customer_dir):
            if name.lower().endswith(".zip"):
                has_zip = True
                break
    if not has_zip:
        missing.append("*.zip (至少需要一个ZIP打包文件)")

    if missing:
        return CheckResult(
            check_name="预期文件齐全",
            passed=False,
            message=f"缺少{len(missing)}个预期文件",
            details="\n".join(missing),
        )

    return CheckResult(
        check_name="预期文件齐全",
        passed=True,
        message="所有预期文件均已生成",
    )


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------
def audit_artifacts(customer_dir: str, identity: str = "", goal: str = "") -> AuditReport:
    """Run all quality checks on customer output directory.

    This is the FINAL quality gate. If any check fails, the pipeline fails.

    Args:
        customer_dir: Absolute path to the customer output directory
                      (e.g. outputs/case_001/)
        identity: User identity string for determining expected files.
        goal: User goal string for goal-specific expected files.

    Returns:
        AuditReport with results of all 8 checks.
    """
    report = AuditReport(case_dir=customer_dir)

    # Validate directory exists
    if not os.path.isdir(customer_dir):
        report.checks.append(
            CheckResult(
                check_name="目录存在",
                passed=False,
                message=f"客户输出目录不存在: {customer_dir}",
            )
        )
        return report

    # Run all 8 checks
    check_functions = [
        _check_docx_openable,          # Check 1
        _check_docx_chinese_content,   # Check 2
        _check_pdf_headers,            # Check 3
        _check_xlsx_rows,              # Check 4
        _check_no_internal_files,      # Check 5
        _check_no_placeholders,        # Check 6
        _check_no_internal_fields,     # Check 7
    ]

    for check_fn in check_functions:
        try:
            result = check_fn(customer_dir)
        except Exception as exc:
            result = CheckResult(
                check_name=check_fn.__name__,
                passed=False,
                message=f"检查执行异常: {exc}",
            )
        report.checks.append(result)

    # Check 8: expected files (identity-aware)
    try:
        result = _check_expected_files(customer_dir, identity=identity, goal=goal)
    except Exception as exc:
        result = CheckResult(
            check_name="预期文件齐全",
            passed=False,
            message=f"检查执行异常: {exc}",
        )
    report.checks.append(result)

    return report
