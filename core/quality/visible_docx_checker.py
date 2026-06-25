"""
visible_docx_checker.py - Quality Gate: DOCX Content Validation

Validates that DOCX files are readable, contain real content (no placeholders),
and do not leak internal system fields to user-facing documents.
"""
from __future__ import annotations

import os
from typing import List

from core.quality.final_artifact_auditor import CheckResult


# Forbidden placeholder text patterns
FORBIDDEN_PLACEHOLDERS: List[str] = [
    "TODO",
    "暂无内容",
    "暂无数据",
    "暂无信息",
    "请自行补充",
    "lorem ipsum",
    "示例文本",
    "待填写",
    "此处填写",
    "XXX",
    "占位符",
]

# Forbidden internal field names that should never appear in user docs
FORBIDDEN_INTERNAL_FIELDS: List[str] = [
    "fact_card",
    "analysis.json",
    "prompt",
    "source_id",
    "confidence",
    "raw_output",
    "workflow_trace",
]


def _read_all_text(docx_path: str) -> str:
    """Extract all visible text from a DOCX file (paragraphs + tables).

    Returns concatenated text or empty string on failure.
    """
    try:
        from docx import Document

        doc = Document(docx_path)
        parts: List[str] = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)

        return "\n".join(parts)
    except Exception:
        return ""


def _count_chinese(text: str) -> int:
    """Count CJK unified ideographs in text."""
    return sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")


def check_docx_readable(docx_path: str) -> CheckResult:
    """Check if DOCX can be opened and has content.

    Validates:
    - File exists and is a real file
    - python-docx can open it
    - Document has at least 1 paragraph
    - Document contains >= 100 Chinese characters

    Args:
        docx_path: Absolute path to the .docx file.

    Returns:
        CheckResult with pass/fail status.
    """
    basename = os.path.basename(docx_path)

    # File existence check
    if not os.path.isfile(docx_path):
        return CheckResult(
            check_name="DOCX可读",
            passed=False,
            message=f"文件不存在: {basename}",
        )

    try:
        from docx import Document

        doc = Document(docx_path)
    except ImportError:
        return CheckResult(
            check_name="DOCX可读",
            passed=False,
            message=f"python-docx未安装，无法打开: {basename}",
        )
    except Exception as exc:
        return CheckResult(
            check_name="DOCX可读",
            passed=False,
            message=f"无法打开DOCX文件: {basename} - {exc}",
        )

    # Count paragraphs (including table cells)
    para_count = len(doc.paragraphs)
    table_text_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_text_count += 1

    total_content = para_count + table_text_count
    if total_content == 0:
        return CheckResult(
            check_name="DOCX可读",
            passed=False,
            message=f"DOCX文件无内容: {basename} (段落数: {para_count})",
        )

    # Check Chinese character count
    full_text = _read_all_text(docx_path)
    cn_count = _count_chinese(full_text)

    if cn_count < 100:
        return CheckResult(
            check_name="DOCX可读",
            passed=False,
            message=f"DOCX中文字符不足: {basename} ({cn_count}个, 需>=100)",
        )

    return CheckResult(
        check_name="DOCX可读",
        passed=True,
        message=f"DOCX可正常打开: {basename} ({para_count}段落, {cn_count}中文字符)",
    )


def check_docx_no_placeholders(docx_path: str) -> CheckResult:
    """Check DOCX has no placeholder text.

    Scans all paragraphs and table cells for forbidden placeholder strings.
    Case-insensitive matching for English patterns, exact matching for Chinese.

    Args:
        docx_path: Absolute path to the .docx file.

    Returns:
        CheckResult with pass/fail status and details of any matches.
    """
    basename = os.path.basename(docx_path)

    if not os.path.isfile(docx_path):
        return CheckResult(
            check_name="无占位文本",
            passed=False,
            message=f"文件不存在: {basename}",
        )

    full_text = _read_all_text(docx_path)
    if not full_text:
        return CheckResult(
            check_name="无占位文本",
            passed=True,
            message=f"DOCX无可读文本: {basename}",
        )

    lines = full_text.split("\n")

    found: List[str] = []
    for placeholder in FORBIDDEN_PLACEHOLDERS:
        placeholder_lower = placeholder.lower()
        for i, line in enumerate(lines, 1):
            if placeholder_lower in line.lower():
                snippet = line.strip()[:80]
                found.append(
                    f"第{i}行发现'{placeholder}': ...{snippet}..."
                )

    if found:
        return CheckResult(
            check_name="无占位文本",
            passed=False,
            message=f"DOCX含{len(found)}处占位文本: {basename}",
            details="\n".join(found),
        )

    return CheckResult(
        check_name="无占位文本",
        passed=True,
        message=f"DOCX无占位文本: {basename}",
    )


def check_docx_no_internal_fields(docx_path: str) -> CheckResult:
    """Check DOCX has no internal system fields.

    Scans all paragraphs and table cells for internal field names that
    should never appear in user-facing documents.

    Args:
        docx_path: Absolute path to the .docx file.

    Returns:
        CheckResult with pass/fail status and details of any matches.
    """
    basename = os.path.basename(docx_path)

    if not os.path.isfile(docx_path):
        return CheckResult(
            check_name="无内部字段",
            passed=False,
            message=f"文件不存在: {basename}",
        )

    full_text = _read_all_text(docx_path)
    if not full_text:
        return CheckResult(
            check_name="无内部字段",
            passed=True,
            message=f"DOCX无可读文本: {basename}",
        )

    lines = full_text.split("\n")

    found: List[str] = []
    for field_name in FORBIDDEN_INTERNAL_FIELDS:
        field_lower = field_name.lower()
        for i, line in enumerate(lines, 1):
            if field_lower in line.lower():
                snippet = line.strip()[:80]
                found.append(
                    f"第{i}行发现内部字段'{field_name}': ...{snippet}..."
                )

    if found:
        return CheckResult(
            check_name="无内部字段",
            passed=False,
            message=f"DOCX含{len(found)}处内部字段: {basename}",
            details="\n".join(found),
        )

    return CheckResult(
        check_name="无内部字段",
        passed=True,
        message=f"DOCX无内部字段泄露: {basename}",
    )
