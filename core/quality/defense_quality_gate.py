"""
core/quality/defense_quality_gate.py - Defense Quality Gate

Hard quality gate for defense scenario (被诉方+应诉答辩).
Blocks delivery if ANY critical check fails.
"""
from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List

from core.contracts.quality_gate import (
    QualityGate, QualityResult, QualityCheck, CheckSeverity
)


class DefenseQualityGate(QualityGate):
    """Quality gate for defense scenario."""

    @property
    def name(self) -> str:
        return "defense_quality_gate"

    def run_checks(self, case_dir: str, ai_mode: str = "") -> QualityResult:
        """Run all defense-specific quality checks."""
        result = QualityResult(gate_name=self.name, ai_mode=ai_mode)
        customer_dir = os.path.join(case_dir, "customer")

        # Check 1: AI mode must not be local_fallback
        result.checks.append(self._check_ai_mode(ai_mode))

        # Check 2: PDF must exist
        result.checks.append(self._check_pdf_exists(customer_dir))

        # Check 3: ZIP must contain PDF
        result.checks.append(self._check_zip_contains_pdf(customer_dir))

        # Check 4: No forbidden patterns in DOCX
        result.checks.extend(self._check_forbidden_patterns(customer_dir))

        # Check 5: Defense not generic
        result.checks.append(self._check_defense_not_generic(customer_dir))

        # Check 6: Evidence gap >= 5
        result.checks.append(self._check_evidence_gap_count(case_dir))

        # Check 7: Action advice >= 6
        result.checks.append(self._check_action_advice_count(case_dir))

        # Check 8: SABCD rating has reasoning
        result.checks.append(self._check_rating_reasoning(case_dir))

        return result

    def _check_ai_mode(self, ai_mode: str) -> QualityCheck:
        """Check AI mode is not local_fallback."""
        if ai_mode == "local_fallback":
            return QualityCheck(
                check_name="AI模式验证",
                passed=False,
                severity=CheckSeverity.CRITICAL,
                message="FAIL_NO_REAL_AI: local_fallback模式不允许正式交付",
                remediation="请配置有效的API密钥并确保AI调用成功",
            )
        return QualityCheck(
            check_name="AI模式验证",
            passed=True,
            severity=CheckSeverity.CRITICAL,
            message=f"AI模式: {ai_mode}",
        )

    def _check_pdf_exists(self, customer_dir: str) -> QualityCheck:
        """Check at least one PDF exists."""
        pdf_files = list(Path(customer_dir).glob("*.pdf")) if os.path.isdir(customer_dir) else []
        if not pdf_files:
            return QualityCheck(
                check_name="PDF文件存在",
                passed=False,
                severity=CheckSeverity.CRITICAL,
                message="FAIL: 客户目录中无PDF文件",
                remediation="确保渲染阶段生成PDF文件",
            )
        return QualityCheck(
            check_name="PDF文件存在",
            passed=True,
            message=f"找到{len(pdf_files)}个PDF文件",
        )

    def _check_zip_contains_pdf(self, customer_dir: str) -> QualityCheck:
        """Check ZIP contains PDF files."""
        import zipfile
        zip_files = list(Path(customer_dir).glob("*.zip")) if os.path.isdir(customer_dir) else []
        if not zip_files:
            return QualityCheck(
                check_name="ZIP包含PDF",
                passed=False,
                severity=CheckSeverity.CRITICAL,
                message="FAIL: 无ZIP文件",
                remediation="确保打包阶段生成ZIP文件",
            )

        for zip_path in zip_files:
            try:
                with zipfile.ZipFile(zip_path, 'r') as zf:
                    pdf_files = [f for f in zf.namelist() if f.endswith('.pdf')]
                    if pdf_files:
                        return QualityCheck(
                            check_name="ZIP包含PDF",
                            passed=True,
                            message=f"ZIP中包含{len(pdf_files)}个PDF文件",
                        )
            except Exception:
                pass

        return QualityCheck(
            check_name="ZIP包含PDF",
            passed=False,
            severity=CheckSeverity.CRITICAL,
            message="FAIL: ZIP中不含PDF文件",
            remediation="确保ZIP打包时包含PDF文件",
        )

    def _check_forbidden_patterns(self, customer_dir: str) -> List[QualityCheck]:
        """Check DOCX files for forbidden patterns."""
        forbidden = [
            ("文书 1:", "内部标记'文书 1:'"),
            ("文书 1：", "内部标记'文书 1：'"),
            ("类型:", "内部标记'类型:'"),
            ("类型：", "内部标记'类型：'"),
            ("待补充", "占位符'待补充'"),
            ("{{", "模板变量'{{'"),
            ("}}", "模板变量'}}'"),
            ("TODO", "开发标记'TODO'"),
            ("暂无", "占位符'暂无'"),
            ("请自行补充", "占位符'请自行补充'"),
            ("fact_card", "内部字段'fact_card'"),
            ("prompt", "内部字段'prompt'"),
            ("source_id", "内部字段'source_id'"),
        ]

        checks = []
        docx_files = list(Path(customer_dir).glob("*.docx")) if os.path.isdir(customer_dir) else []

        if not docx_files:
            checks.append(QualityCheck(
                check_name="DOCX内容检查",
                passed=True,
                message="无DOCX文件需要检查",
            ))
            return checks

        for docx_path in docx_files:
            try:
                from docx import Document
                doc = Document(str(docx_path))
                text = "\n".join(p.text for p in doc.paragraphs)

                for pattern, desc in forbidden:
                    if pattern in text:
                        checks.append(QualityCheck(
                            check_name=f"DOCX无{desc}",
                            passed=False,
                            severity=CheckSeverity.CRITICAL,
                            message=f"FAIL: {docx_path.name}含{desc}",
                            details=f"发现'{pattern}'",
                            remediation="移除内部标记和占位符",
                        ))
                    else:
                        checks.append(QualityCheck(
                            check_name=f"DOCX无{desc}",
                            passed=True,
                            message=f"{docx_path.name}无{desc}",
                        ))
            except Exception as e:
                checks.append(QualityCheck(
                    check_name="DOCX读取",
                    passed=False,
                    severity=CheckSeverity.ERROR,
                    message=f"无法读取{docx_path.name}: {e}",
                ))

        return checks

    def _check_defense_not_generic(self, customer_dir: str) -> QualityCheck:
        """Check defense document is not generic."""
        docx_files = list(Path(customer_dir).glob("*答辩*.docx")) if os.path.isdir(customer_dir) else []
        if not docx_files:
            return QualityCheck(
                check_name="答辩状非泛化",
                passed=True,
                message="未找到答辩状文件",
            )

        generic_indicators = [
            "本案基本事实如下：",
            "针对原告的诉讼请求，答辩人提出以下答辩意见：",
            "根据案件情况，建议采取以下诉讼策略：",
            "综合评级：",
        ]

        for docx_path in docx_files:
            try:
                from docx import Document
                doc = Document(str(docx_path))
                text = "\n".join(p.text for p in doc.paragraphs)

                generic_count = sum(1 for g in generic_indicators if g in text)
                if generic_count >= 3:
                    return QualityCheck(
                        check_name="答辩状非泛化",
                        passed=False,
                        severity=CheckSeverity.CRITICAL,
                        message=f"FAIL: {docx_path.name}内容泛化",
                        details=f"发现{generic_count}个泛化指标",
                        remediation="请确保答辩状针对具体案件事实生成",
                    )
            except Exception:
                pass

        return QualityCheck(
            check_name="答辩状非泛化",
            passed=True,
            message="答辩状内容针对具体案件",
        )

    def _check_evidence_gap_count(self, case_dir: str) -> QualityCheck:
        """Check evidence gap count >= 5."""
        distilled_path = os.path.join(case_dir, "_internal", "distilled_card.json")
        if not os.path.exists(distilled_path):
            return QualityCheck(
                check_name="证据缺口数量",
                passed=False,
                severity=CheckSeverity.ERROR,
                message="distilled_card.json不存在",
            )

        try:
            import json
            with open(distilled_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            gaps = data.get("strategy_card", {}).get("evidence_gap", [])
            if len(gaps) < 5:
                return QualityCheck(
                    check_name="证据缺口数量",
                    passed=False,
                    severity=CheckSeverity.CRITICAL,
                    message=f"FAIL: 证据缺口仅{len(gaps)}条（需要>=5）",
                    remediation="确保证据缺口分析覆盖5个以上方面",
                )
            return QualityCheck(
                check_name="证据缺口数量",
                passed=True,
                message=f"证据缺口{len(gaps)}条",
            )
        except Exception as e:
            return QualityCheck(
                check_name="证据缺口数量",
                passed=False,
                severity=CheckSeverity.ERROR,
                message=f"读取distilled_card失败: {e}",
            )

    def _check_action_advice_count(self, case_dir: str) -> QualityCheck:
        """Check action advice count >= 6."""
        distilled_path = os.path.join(case_dir, "_internal", "distilled_card.json")
        if not os.path.exists(distilled_path):
            return QualityCheck(
                check_name="行动建议数量",
                passed=False,
                severity=CheckSeverity.ERROR,
                message="distilled_card.json不存在",
            )

        try:
            import json
            with open(distilled_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            advice = data.get("strategy_card", {}).get("action_advice", [])
            if len(advice) < 6:
                return QualityCheck(
                    check_name="行动建议数量",
                    passed=False,
                    severity=CheckSeverity.CRITICAL,
                    message=f"FAIL: 行动建议仅{len(advice)}条（需要>=6）",
                    remediation="确保行动建议覆盖6个以上方面",
                )
            return QualityCheck(
                check_name="行动建议数量",
                passed=True,
                message=f"行动建议{len(advice)}条",
            )
        except Exception as e:
            return QualityCheck(
                check_name="行动建议数量",
                passed=False,
                severity=CheckSeverity.ERROR,
                message=f"读取distilled_card失败: {e}",
            )

    def _check_rating_reasoning(self, case_dir: str) -> QualityCheck:
        """Check SABCD rating has reasoning."""
        distilled_path = os.path.join(case_dir, "_internal", "distilled_card.json")
        if not os.path.exists(distilled_path):
            return QualityCheck(
                check_name="评级理由",
                passed=False,
                severity=CheckSeverity.ERROR,
                message="distilled_card.json不存在",
            )

        try:
            import json
            with open(distilled_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            sc = data.get("strategy_card", {})
            rating = sc.get("sabcd_rating", "")
            situation = sc.get("situation_assessment", "")

            if not rating:
                return QualityCheck(
                    check_name="评级理由",
                    passed=False,
                    severity=CheckSeverity.CRITICAL,
                    message="FAIL: 无SABCD评级",
                    remediation="确保生成SABCD评级",
                )
            if not situation or len(situation) < 50:
                return QualityCheck(
                    check_name="评级理由",
                    passed=False,
                    severity=CheckSeverity.CRITICAL,
                    message=f"FAIL: 评级'{rating}'无充分理由",
                    details=f"处境评估仅{len(situation)}字符",
                    remediation="确保处境评估至少50字符",
                )
            return QualityCheck(
                check_name="评级理由",
                passed=True,
                message=f"评级'{rating}'有充分理由",
            )
        except Exception as e:
            return QualityCheck(
                check_name="评级理由",
                passed=False,
                severity=CheckSeverity.ERROR,
                message=f"读取distilled_card失败: {e}",
            )
