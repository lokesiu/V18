"""
step8_quality_gate.py - Pipeline Step 8: Quality Gate / Artifact Audit

Runs comprehensive quality checks on all rendered artifacts in the
customer delivery directory. Validates file integrity, content quality,
and format correctness.

Final pipeline step - determines whether the output PASSes or FAILs.
"""
from __future__ import annotations

import os
from typing import List, Tuple

from core.fact_card import PipelineContext
from core.quality.final_artifact_auditor import audit_artifacts


# Quality check thresholds
MIN_XLSX_ROWS = 3
MIN_PDF_HEADER = b"%PDF"
MAX_PLACEHOLDER_WARNINGS = 5


def _check_docx_files(customer_dir: str, ctx: PipelineContext) -> Tuple[bool, List[str]]:
    """Verify all DOCX files open correctly and contain Chinese text.

    Args:
        customer_dir: Directory containing rendered artifacts.
        ctx: PipelineContext for logging.

    Returns:
        Tuple of (all_passed, list_of_issues).
    """
    issues = []

    try:
        from docx import Document
    except ImportError:
        issues.append("python-docx 未安装，无法验证 DOCX 文件")
        return False, issues

    for filename in os.listdir(customer_dir):
        if not filename.endswith(".docx"):
            continue

        filepath = os.path.join(customer_dir, filename)

        try:
            doc = Document(filepath)
            # Check file is not empty
            if len(doc.paragraphs) == 0:
                issues.append(f"{filename}: DOCX 文件无段落内容")
                continue

            # Check for Chinese text content
            full_text = "\n".join(p.text for p in doc.paragraphs)
            has_chinese = any("\u4e00" <= ch <= "\u9fff" for ch in full_text)

            if not has_chinese:
                issues.append(f"{filename}: DOCX 文件不含中文文本内容")

            # Check for TODO/placeholder text
            # Note: "待补充" is a legitimate status indicator for missing evidence
            placeholder_patterns = [
                "TODO", "FIXME", "XXX", "PLACEHOLDER",
                "请自行补充", "lorem ipsum", "示例文本",
            ]
            for pattern in placeholder_patterns:
                if pattern in full_text:
                    issues.append(f"{filename}: DOCX 包含占位符文本 '{pattern}'")

            # Check for internal/debug fields
            internal_patterns = [
                "DEBUG", "INTERNAL", "SECRET", "PASSWORD",
                "API_KEY", "TOKEN",
            ]
            for pattern in internal_patterns:
                if pattern.upper() in full_text.upper():
                    issues.append(f"{filename}: DOCX 包含内部字段 '{pattern}'")

        except Exception as exc:
            issues.append(f"{filename}: DOCX 文件无法打开 - {exc}")

    return len(issues) == 0, issues


def _check_pdf_files(customer_dir: str, ctx: PipelineContext) -> Tuple[bool, List[str]]:
    """Verify all PDF files have valid %PDF header.

    Args:
        customer_dir: Directory containing rendered artifacts.
        ctx: PipelineContext for logging.

    Returns:
        Tuple of (all_passed, list_of_issues).
    """
    issues = []

    for filename in os.listdir(customer_dir):
        if not filename.endswith(".pdf"):
            continue

        filepath = os.path.join(customer_dir, filename)

        try:
            with open(filepath, "rb") as f:
                header = f.read(8)

            if not header.startswith(MIN_PDF_HEADER):
                issues.append(f"{filename}: PDF 文件头无效 (期望 %PDF，实际 {header[:8]})")

            # Check file size is reasonable (not empty/tiny)
            file_size = os.path.getsize(filepath)
            if file_size < 100:
                issues.append(f"{filename}: PDF 文件过小 ({file_size} bytes)，可能损坏")

        except Exception as exc:
            issues.append(f"{filename}: PDF 文件验证失败 - {exc}")

    return len(issues) == 0, issues


def _check_xlsx_files(customer_dir: str, ctx: PipelineContext) -> Tuple[bool, List[str]]:
    """Verify all XLSX files have >= 3 rows.

    Args:
        customer_dir: Directory containing rendered artifacts.
        ctx: PipelineContext for logging.

    Returns:
        Tuple of (all_passed, list_of_issues).
    """
    issues = []

    try:
        from openpyxl import load_workbook
    except ImportError:
        issues.append("openpyxl 未安装，无法验证 XLSX 文件")
        return False, issues

    for filename in os.listdir(customer_dir):
        if not filename.endswith(".xlsx"):
            continue

        filepath = os.path.join(customer_dir, filename)

        try:
            wb = load_workbook(filepath, read_only=True)
            total_rows = 0
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                row_count = ws.max_row or 0
                total_rows += row_count
            wb.close()

            if total_rows < MIN_XLSX_ROWS:
                issues.append(
                    f"{filename}: XLSX 仅 {total_rows} 行 (要求 >= {MIN_XLSX_ROWS})"
                )

        except Exception as exc:
            issues.append(f"{filename}: XLSX 文件验证失败 - {exc}")

    return len(issues) == 0, issues


def _check_no_debug_files(customer_dir: str, ctx: PipelineContext) -> Tuple[bool, List[str]]:
    """Verify no JSON/MD/debug files leaked into customer directory.

    Args:
        customer_dir: Directory containing rendered artifacts.
        ctx: PipelineContext for logging.

    Returns:
        Tuple of (all_passed, list_of_issues).
    """
    issues = []
    forbidden_extensions = {".json", ".md", ".py", ".log", ".debug", ".tmp"}

    for filename in os.listdir(customer_dir):
        _, ext = os.path.splitext(filename)
        if ext.lower() in forbidden_extensions:
            issues.append(
                f"{filename}: 客户交付目录不应包含 {ext} 类型文件"
            )

    # Check for hidden files
    for filename in os.listdir(customer_dir):
        if filename.startswith("."):
            issues.append(f"{filename}: 客户交付目录不应包含隐藏文件")

    return len(issues) == 0, issues


def step8_quality_gate(ctx: PipelineContext) -> PipelineContext:
    """Run comprehensive quality audit on all rendered artifacts.

    Checks performed:
    1. DOCX files open, contain Chinese text, no placeholders
    2. PDF files have valid %PDF header
    3. XLSX files have >= 3 rows
    4. No JSON/MD/debug files in customer directory
    5. No TODO/placeholder text
    6. No internal/debug fields leaked
    7. Final artifact auditor validation

    Args:
        ctx: PipelineContext with rendered files in output_dir/customer/.

    Returns:
        PipelineContext with errors list updated (empty = PASS).
    """
    ctx.log("Step 8: 质量门禁 - 对交付物进行全面质量审计")

    customer_dir = os.path.join(ctx.output_dir, "customer")

    if not os.path.isdir(customer_dir):
        ctx.add_error(f"客户交付目录不存在: {customer_dir}")
        return ctx

    # List all files in customer directory
    all_files = os.listdir(customer_dir)
    ctx.log(f"  客户交付目录包含 {len(all_files)} 个文件")

    for filename in sorted(all_files):
        filepath = os.path.join(customer_dir, filename)
        size = os.path.getsize(filepath)
        ctx.log(f"  - {filename} ({size:,} bytes)")

    total_issues: List[str] = []

    # --- Run all quality checks ---

    # 1. DOCX validation
    ctx.log("  检查 DOCX 文件...")
    docx_ok, docx_issues = _check_docx_files(customer_dir, ctx)
    total_issues.extend(docx_issues)
    if docx_ok:
        ctx.log("  DOCX 检查通过")
    else:
        ctx.log(f"  DOCX 检查发现 {len(docx_issues)} 个问题")

    # 2. PDF validation
    ctx.log("  检查 PDF 文件...")
    pdf_ok, pdf_issues = _check_pdf_files(customer_dir, ctx)
    total_issues.extend(pdf_issues)
    if pdf_ok:
        ctx.log("  PDF 检查通过")
    else:
        ctx.log(f"  PDF 检查发现 {len(pdf_issues)} 个问题")

    # 3. XLSX validation
    ctx.log("  检查 XLSX 文件...")
    xlsx_ok, xlsx_issues = _check_xlsx_files(customer_dir, ctx)
    total_issues.extend(xlsx_issues)
    if xlsx_ok:
        ctx.log("  XLSX 检查通过")
    else:
        ctx.log(f"  XLSX 检查发现 {len(xlsx_issues)} 个问题")

    # 4. Forbidden file types
    ctx.log("  检查禁止文件类型...")
    debug_ok, debug_issues = _check_no_debug_files(customer_dir, ctx)
    total_issues.extend(debug_issues)
    if debug_ok:
        ctx.log("  文件类型检查通过")
    else:
        ctx.log(f"  文件类型检查发现 {len(debug_issues)} 个问题")

    # 5. Run final artifact auditor
    ctx.log("  运行最终交付物审计...")
    try:
        from core.quality.final_artifact_auditor import AuditReport
        audit_result = audit_artifacts(customer_dir, identity=ctx.identity)
        if audit_result is not None:
            # audit_artifacts returns AuditReport object, not dict
            if isinstance(audit_result, AuditReport):
                audit_passed = audit_result.passed
                if not audit_passed:
                    for check in audit_result.failed_checks:
                        issue = f"[审计] {check.check_name}: {check.message}"
                        if check.details:
                            issue += f" | {check.details}"
                        total_issues.append(issue)
                ctx.log(f"  最终审计: {audit_result.summary()}")
                for check in audit_result.checks:
                    status = "[OK]" if check.passed else "[FAIL]"
                    ctx.log(f"    {status} {check.check_name}: {check.message}")
            else:
                # Fallback for unexpected return type
                ctx.log(f"  最终审计: 返回类型异常 ({type(audit_result).__name__})")
                total_issues.append(f"审计返回类型异常: {type(audit_result).__name__}")
    except ImportError:
        ctx.log("  WARNING: final_artifact_auditor 模块不可用，跳过最终审计")
        total_issues.append("final_artifact_auditor 模块不可用")
    except Exception as exc:
        ctx.log(f"  WARNING: 最终审计异常: {exc}")
        total_issues.append(f"最终审计异常: {exc}")

    # --- Final verdict ---
    if total_issues:
        for issue in total_issues:
            ctx.add_error(f"[质量门禁] {issue}")
        ctx.log(
            f"Step 8 完成: 质量门禁 FAIL - 发现 {len(total_issues)} 个问题"
        )
    else:
        ctx.log("Step 8 完成: 质量门禁 PASS - 所有检查通过")

    # Log summary statistics
    file_counts = {"docx": 0, "pdf": 0, "xlsx": 0, "zip": 0, "other": 0}
    for filename in all_files:
        _, ext = os.path.splitext(filename)
        ext_lower = ext.lower().lstrip(".")
        if ext_lower in file_counts:
            file_counts[ext_lower] += 1
        else:
            file_counts["other"] += 1

    ctx.log(
        f"  交付物统计: {file_counts['docx']} DOCX + "
        f"{file_counts['pdf']} PDF + "
        f"{file_counts['xlsx']} XLSX + "
        f"{file_counts['zip']} ZIP + "
        f"{file_counts['other']} 其他"
    )

    # Save AI mode manifest
    tracker = getattr(ctx, '_ai_mode_tracker', None)
    if tracker:
        tracker.finish()
        tracker.save_manifest(ctx.output_dir)
        ctx.log(f"  AI 模式: {tracker.ai_mode.value}")
        ctx.log(f"  API-A 状态: {tracker.api_a_status.value} ({tracker.api_a_latency_ms}ms)")
        ctx.log(f"  API-B 状态: {tracker.api_b_status.value} ({tracker.api_b_latency_ms}ms)")

    return ctx
